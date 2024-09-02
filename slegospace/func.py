from sklearn.model_selection import train_test_split
import pandas as pd
from typing import Union
import joblib
# Import necessary libraries

def df_keep_rows_by_index(input_csv_file: str = 'dataspace/dataset.csv',
                       start_index: int = 0,
                       end_index: int = 100,
                       output_csv_file: str = 'dataspace/dataset_selected_rows.csv'):
    """
    Keeps rows in a pandas DataFrame based on specified index range.

    Parameters:
    input_csv_file (str): The file path to read the CSV file. Default is 'dataspace/dataset.csv'.
    start_index (int): The start index of the row range to keep.
    end_index (int): The end index of the row range to keep. This index is inclusive.
    output_csv_file (str): The file path to save the modified CSV file. Default is 'dataspace/selected_rows_dataset.csv'.

    Returns:
    pd.DataFrame: The DataFrame containing only the rows within the specified index range.
    """
    # Read the CSV file into a DataFrame
    df = pd.read_csv(input_csv_file)
    
    # Validate index range
    if start_index < 0 or end_index >= len(df):
        raise ValueError("Start or end index is out of the DataFrame's range.")
    
    # Keep only the rows within the specified index range
    df = df.iloc[start_index:end_index + 1]  # +1 because the end index is inclusive
    
    # Save the modified DataFrame to a CSV file
    df.to_csv(output_csv_file, index=False)
    
    return df


def df_keep_columns(input_csv_file: str = 'dataspace/dataset.csv',
                          columns_to_keep: list = ['ImportantColumn1', 'ImportantColumn2'],
                          output_csv_file: str = 'dataspace/dataset_filtered.csv'):
    """
    Keeps specified columns in a pandas DataFrame and optionally filters rows based on a condition.

    Parameters:
    input_csv_file (str): The file path to read the CSV file. Default is 'dataspace/dataset.csv'.
    columns_to_keep (list): List of column names to be retained in the DataFrame.
                            Default is ['ImportantColumn1', 'ImportantColumn2'].
    output_csv_file (str): The file path to save the modified CSV file. Default is 'dataspace/filtered_dataset.csv'.

    Returns:
    pd.DataFrame: The DataFrame with only the specified columns and optionally filtered rows.
    """
    # Read the CSV file into a DataFrame
    df = pd.read_csv(input_csv_file)
    
    # Check if all specified columns to keep are present in the DataFrame
    missing_columns = [col for col in columns_to_keep if col not in df.columns]
    if missing_columns:
        raise ValueError(f"Columns not found in DataFrame: {missing_columns}")
    
    # Keep only the specified columns
    df = df[columns_to_keep]
    
    # Save the modified DataFrame to a CSV file
    df.to_csv(output_csv_file, index=False)
    
    return df



def df_delete_columns(input_csv_file: str = 'dataspace/dataset.csv',
                   columns_to_delete: list = ['UnwantedColumn'],
                   output_csv_file: str = 'dataspace/dataset_cleaned.csv'):
    """
    Deletes specified columns from a pandas DataFrame.

    Parameters:
    input_csv_file (str): The file path to read the CSV file. Default is 'dataspace/dataset.csv'.
    columns_to_delete (list): List of column names to be deleted from the DataFrame.
                              Default is ['UnwantedColumn'].
    output_csv_file (str): The file path to save the modified CSV file. Default is 'dataspace/cleaned_dataset.csv'.

    Returns:
    pd.DataFrame: The DataFrame with the specified columns removed.
    """
    # Read the CSV file into a DataFrame
    df = pd.read_csv(input_csv_file)
    
    # Check if all specified columns to delete are present in the DataFrame
    missing_columns = [col for col in columns_to_delete if col not in df.columns]
    if missing_columns:
        raise ValueError(f"Columns not found in DataFrame: {missing_columns}")
    
    # Delete the specified columns
    df.drop(columns=columns_to_delete, inplace=True)
    
    # Save the modified DataFrame to a CSV file
    df.to_csv(output_csv_file, index=False)
    
    return df

def df_rename_columns(input_csv_file: str = 'dataspace/dataset.csv',
                      rename_dict: dict = {'Close': 'close', 'Open': 'open'},
                      output_csv_file: str = 'dataspace/dataset_renamed.csv'):
    """
    Renames multiple columns in a pandas DataFrame based on a dictionary mapping from old names to new names.

    Parameters:
    input_csv_file (str): The file path to read the CSV file. Default is 'dataspace/dataset.csv'.
    rename_dict (dict): Dictionary where keys are old column names and values are new column names.
                        Default is {'Close': 'close', 'Open': 'open'}.
    output_csv_file (str): The file path to save the modified CSV file. Default is 'dataspace/modified_dataset.csv'.

    Returns:
    pd.DataFrame: The DataFrame with renamed columns.
    """
    # Read the CSV file into a DataFrame
    df = pd.read_csv(input_csv_file)
    
    # Check if all keys in the dictionary are present in the DataFrame columns
    missing_columns = [col for col in rename_dict if col not in df.columns]
    if missing_columns:
        raise ValueError(f"Columns not found in DataFrame: {missing_columns}")
    
    # Rename the columns
    df.rename(columns=rename_dict, inplace=True)
    
    # Save the modified DataFrame to a CSV file
    df.to_csv(output_csv_file, index=False)
    
    return df

def merge_csv_by_index(input_csv_file1: str = 'dataspace/dataset1.csv',
                       input_csv_file2: str = 'dataspace/dataset2.csv',
                       output_csv_file: str = 'dataspace/merged_dataset.csv',
                       rename_dict1: dict = None,
                       rename_dict2: dict = None,
                       how: str = 'outer') -> pd.DataFrame:
    """
    Merges two CSV files horizontally based on their index after optionally renaming columns in each.

    Parameters:
    input_csv_file1 (str): Path to the first CSV file. Default path is provided.
    input_csv_file2 (str): Path to the second CSV file. Default path is provided.
    output_csv_file (str): Path where the merged CSV file will be saved. Default path is provided.
    rename_dict1 (dict): Dictionary to rename columns of the first CSV file.
    rename_dict2 (dict): Dictionary to rename columns of the second CSV file.
    how (str): Type of merge to perform ('outer', 'inner', 'left', 'right'). Default is 'outer'.

    Returns:
    pd.DataFrame: A dataframe containing the merged data.
    """
    # Read the CSV files into DataFrames
    df1 = pd.read_csv(input_csv_file1)
    df2 = pd.read_csv(input_csv_file2)

    # Rename columns if dictionaries are provided
    if rename_dict1:
        df1.rename(columns=rename_dict1, inplace=True)
    if rename_dict2:
        df2.rename(columns=rename_dict2, inplace=True)

    # Merge the DataFrames based on the index
    merged_df = pd.merge(df1, df2, left_index=True, right_index=True, how=how)
    # Drop any 'Unnamed' columns that may exist
    merged_df = merged_df.loc[:, ~merged_df.columns.str.contains('^Unnamed')]

    # Save the merged DataFrame to a CSV file
    merged_df.to_csv(output_csv_file, index=False)

    return merged_df




def prepare_dataset_tabular_ml(input_file_path: str = 'dataspace/AirQuality.csv', 
                               index_col: Union[int, bool] = 0,
                               target_column_name: str = 'NO2(GT)', 
                               drop_columns: list = ['Date','NO2(GT)','Time'],
                               output_file_path:str = 'dataspace/prepared_dataset_tabular_ml.csv'):
    """
    Prepares a dataset for tabular machine learning by loading, cleaning, and optionally modifying it.
    """
    
    # Load the dataset
    df = pd.read_csv(input_file_path, index_col=index_col)
    
    # Drop any 'Unnamed' columns that may exist
    df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
    
    # Set the target column if it exists
    if target_column_name in df.columns:
        df['target'] = df[target_column_name].values
    
    # Drop specified columns if any
    if drop_columns is not None:
        df = df.drop(columns=drop_columns, errors='ignore')
    
    # Save the cleaned dataset
    output_file_path = output_file_path
    df.to_csv(output_file_path)
    
    return df



def split_dataset_4ml(input_data_file:str = 'dataspace/prepared_dataset_tabular_ml.csv', 
                           index_col: Union[int, bool] = 0,
                           train_size:int =0.6, 
                           val_size:int =0.2, 
                           test_size:int =0.2,
                           output_train_file:str ='dataspace/train.csv', 
                           output_val_file:str ='dataspace/val.csv', 
                           output_test_file:str ='dataspace/test.csv'):
    """
    Split a dataset into training, validation, and test sets and save them as CSV files.
    
    Parameters:
    - data (DataFrame): The dataset to split.
    - train_size (float): The proportion of the dataset to include in the train split.
    - val_size (float): The proportion of the dataset to include in the validation split.
    - test_size (float): The proportion of the dataset to include in the test split.
    - train_path (str): File path to save the training set CSV.
    - val_path (str): File path to save the validation set CSV.
    - test_path (str): File path to save the test set CSV.
    
    Returns:
    - train_data (DataFrame): The training set.
    - val_data (DataFrame): The validation set.
    - test_data (DataFrame): The test set.
    """

    # Load data
    data = pd.read_csv(input_data_file,index_col=index_col)

    if not (0 < train_size < 1) or not (0 < val_size < 1) or not (0 < test_size < 1):
        raise ValueError("All size parameters must be between 0 and 1.")
        
    if train_size + val_size + test_size != 1:
        raise ValueError("The sum of train_size, val_size, and test_size must equal 1.")
        
    # First split to separate out the test set
    temp_train_data, test_data = train_test_split(data, test_size=test_size, random_state=42)
    
    # Adjust val_size proportion to account for the initial split (for the remaining data)
    adjusted_val_size = val_size / (1 - test_size)
    
    # Second split to separate out the validation set from the temporary training set
    train_data, val_data = train_test_split(temp_train_data, test_size=adjusted_val_size, random_state=42)
    
    # Save to CSV
    train_data.to_csv(output_train_file, index=False)
    val_data.to_csv(output_val_file, index=False)
    test_data.to_csv(output_test_file, index=False)
    
    return train_data, val_data, test_data


def df_shift_data_row(input_csv_file: str = 'dataspace/dataset.csv',
                      shift_columns: list = ['Close'],
                      shift_columns_name: list = ['Close_shifted'],
                      shift_columns_keep: bool = False,
                      shift_rows: int = -1,
                      keep_row: bool = False,
                      output_csv_file: str = 'dataspace/dataset_selected_rows.csv'):
    """
    Shifts the data in specified columns by the specified number of rows and optionally renames and retains original columns.

    Args:
    input_csv_file (str): Path to the input CSV file.
    shift_columns (list): List of column names whose data will be shifted.
    shift_columns_name (list): List of new names for the shifted columns.
    shift_columns_keep (bool): If True, keeps the original columns alongside the shifted ones.
    shift_rows (int): Number of rows to shift the data by. 1 is one day delay, -1 is one day ahead
    keep_row (bool): If False, rows resulting in NaN values from the shift will be dropped.
    output_csv_file (str): Path to save the output CSV file after shifting the data.
    
    Returns:
    pd.DataFrame: DataFrame with the data shifted and potentially renamed in specified columns.
    """
    # Read the CSV file into a DataFrame
    df = pd.read_csv(input_csv_file)

    # Shift the data in the specified columns and optionally rename
    for orig, new in zip(shift_columns, shift_columns_name):
        df[new] = df[orig].shift(periods=shift_rows)
        if not shift_columns_keep:
            df.drop(orig, axis=1, inplace=True)

    # If keep_row is False, drop rows with NaN values that result from shifting
    if not keep_row:
        df = df.dropna(subset=shift_columns_name)

    # Save the modified DataFrame to a new CSV file
    df.to_csv(output_csv_file, index=False)

    return df










from ucimlrepo import fetch_ucirepo 
def fetch_uci_dataset(uci_data_id:int=360, 
                      output_file_path:str='dataspace/AirQuality.csv'):


    # fetch dataset
    downloaded_data = fetch_ucirepo(id=uci_data_id)

    # data (as pandas dataframes)
    X = downloaded_data.data.features
    y = downloaded_data.data.targets

    # # metadata
    # print(air_quality.metadata)

    # # variable information
    # print(air_quality.variables)
    

    # Assuming `X` and `y` are both pandas DataFrames with the same index
    df = pd.concat([X, y], axis=1)
    df.to_csv(output_file_path)
    return df

import os
import ast
import panel as pn
import inspect
import json
def test_function(input_string:str='Hello!', 
          output_file_path:str='dataspace/output.txt'):
    """
    A simple function to save the provided input string to a specified text file and return the string.

    Parameters:
    - input_string (str): The string to be saved.
    - output_file_path (str): The file path where the string should be saved.

    Returns:
    - str: The same input string.
    """

    # Open the file at the specified path in write mode
    with open(output_file_path, 'w') as file:
        # Write the input string to the file
        file.write(input_string)

    # Return the input string
    return input_string

def _compute(module_name, input):
    module = __import__(module_name)

    #pipeline_dict = ast.literal_eval(input)
    pipeline_dict = json.loads(text)
    output = ""
    for function_name, parameters in pipeline_dict.items():
        function = eval(f"module.{function_name}")
        result = function(**parameters)

        output += "\n===================="+function_name+"====================\n\n"
        output += str(result)

    return output

def _create_multi_select_combobox(target_module):
    """
    Creates a multi-select combobox with all functions from the target_module.
    """
    
    # Get the module name (e.g., "func" if your module is named func.py)
    module_name = target_module.__name__
    
    # Get a list of all functions defined in target_module
    functions = [name for name, obj in inspect.getmembers(target_module, inspect.isfunction)
                 if obj.__module__ == module_name and not name.startswith('_')]

    # Create a multi-select combobox using the list of functions
    multi_combobox = pn.widgets.MultiChoice(name='Functions:', options=functions, height=150)

    return multi_combobox


# def _create_multi_select_combobox(func):
#   """
#   Creates a multi-select combobox with all functions from the func.py file.
#   """

#   # Get a list of all functions in the func.py file.
#   functions = [name for name, obj in inspect.getmembers(func)
#                 if inspect.isfunction(obj) and not name.startswith('_')]

#   # Create a multi-select combobox using the list of functions.
#   multi_combobox = pn.widgets.MultiChoice(name='Functions:', options=functions,  height=150)

#   return multi_combobox


def _extract_parameter(func):
    """
    Extracts the names and default values of the parameters of a function as a dictionary.

    Args:
        func: The function to extract parameter names and default values from.

    Returns:
        A dictionary where the keys are parameter names and the values are the default values.
    """
    signature = inspect.signature(func)
    parameters = signature.parameters

    parameter_dict = {}
    for name, param in parameters.items():
        if param.default != inspect.Parameter.empty:
            parameter_dict[name] = param.default
        else:
            parameter_dict[name] = None

    return parameter_dict


import os
import json

def __combine_json_files(directory, output_file):
    """
    Combine all JSON files in a directory into a single JSON file.

    Args:
    directory (str): The directory containing JSON files.
    output_file (str): The path to the output JSON file.
    """
    combined_data = []  # List to hold data from all JSON files

    # Loop over all files in the directory
    for filename in os.listdir(directory):
        if filename.endswith('.json'):  # Check for JSON files
            file_path = os.path.join(directory, filename)
            with open(file_path, 'r') as file:
                data = json.load(file)  # Load JSON data from file
                combined_data.append(data)  # Append data to the list

    # Write combined data to output JSON file
    with open(output_file, 'w') as file:
        json.dump(combined_data, file, indent=4)  # Use 'indent' for pretty-printing

    print("All JSON files have been combined into:", output_file)

# # Example usage:
# combine_json_files('/content/drive/MyDrive/SLEGO/slegospace/knowledgespace', '/content/drive/MyDrive/SLEGO/slegospace/knowledgespace/kownledge.json')

import requests
from bs4 import BeautifulSoup

def webscrape_to_txt(url: str = "https://au.finance.yahoo.com/", 
                            output_filename: str = "dataspace/output_webscrape.txt"):
    """
    Fetches the content from the specified URL and saves the textual content 
    into a text file, stripping out all HTML tags.

    Parameters:
    - url (str): The URL from which to fetch the content. Default is Yahoo Finance homepage.
    - output_filename (str): The path to the file where the text content will be saved.

    Returns:
    - None: Outputs a file with the extracted text content.
    """
    try:
        # Send a HTTP request to the URL
        response = requests.get(url)
        # Check if the request was successful
        if response.status_code == 200:
            # Parse the HTML content
            soup = BeautifulSoup(response.text, 'html.parser')
            # Extract text using .get_text()
            text_content = soup.get_text(separator='\n', strip=True)
            # Open a file in write mode
            with open(output_filename, 'w', encoding='utf-8') as file:
                file.write(text_content)
            print(f"Text content saved successfully to {output_filename}")
        else:
            print(f"Failed to retrieve the webpage. Status code: {response.status_code}")

        return text_content
    except Exception as e:
        print(f"An error occurred: {e}")


import json
import openai
import os
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import pytesseract
import csv
import json
import openai
import os
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import pytesseract
import csv
from openai import OpenAI

def chatbot_huggingface_api(API_URL: str = "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.2",
                            api_key: str = "your_api_key",
                            user_input_file: str = '',
                            output_text_file: str = 'dataspace/gpt_output_text.txt',
                            output_json_file: str = 'dataspace/gpt_output_full.json',
                            query: dict = {"inputs": "Can you please let us know more details about your "}):
    """
    Sends a query, optionally augmented with contents from various file types, to the specified Hugging Face API endpoint.
    
    This function supports processing inputs from text, PDF, DOCX, XLSX, image files (for OCR), JSON, and CSV files. The
    contents of the file are appended to a base query provided in the `query` parameter. If a file is specified but does
    not exist, the function will raise a FileNotFoundError. Unsupported file types will raise a ValueError.

    Parameters:
        API_URL (str): The URL of the Hugging Face API model to which the request will be sent.
        api_key (str): The API key required for authentication with the Hugging Face API.
        user_input_file (str): The path to the file whose contents are to be appended to the query. If empty, only the base query is used.
        output_text_file (str): The path where the plain text part of the response will be saved.
        output_json_file (str): The path where the full JSON response from the API will be saved.
        query (dict): A dictionary containing the base query. Expected to have at least a key 'inputs' with a starting string.

    Returns:
        dict: A dictionary containing the JSON response from the API. If there's an HTTP error, it returns a dictionary
              containing the error message and status code.

    Raises:
        FileNotFoundError: If the specified `user_input_file` does not exist.
        ValueError: If the file extension of `user_input_file` is not supported.

    Example Usage:
        result = chatbot_huggingface_api(
            api_key="your_api_key_here",
            user_input_file="path/to/input.txt",
            query={"inputs": "Please analyze the following data: "}
        )
        print(result)
    """
    # Function implementation here
    # Initialize the combined message with the query's input
    combined_message = query["inputs"]

    # Process file if specified and exists
    if user_input_file and os.path.exists(user_input_file):
        file_extension = user_input_file.split('.')[-1].lower()

        if file_extension == 'txt':
            with open(user_input_file, 'r') as file:
                file_contents = file.read().strip()
            combined_message += f"\n\n==== Text File Input ====\n\n{file_contents}"
        elif file_extension == 'pdf':
            with open(user_input_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pdf_contents = ' '.join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            combined_message += f"\n\n==== PDF File Input ====\n\n{pdf_contents}"
        elif file_extension == 'docx':
            doc = Document(user_input_file)
            docx_contents = ' '.join([para.text for para in doc.paragraphs])
            combined_message += f"\n\n==== DOCX File Input ====\n\n{docx_contents}"
        elif file_extension == 'xlsx':
            workbook = load_workbook(filename=user_input_file)
            sheet = workbook.active
            xlsx_contents = ' '.join([str(cell.value) for row in sheet for cell in row if cell.value is not None])
            combined_message += f"\n\n==== XLSX File Input ====\n\n{xlsx_contents}"
        elif file_extension in ['png', 'jpg', 'jpeg']:
            img = Image.open(user_input_file)
            image_text = pytesseract.image_to_string(img)
            combined_message += f"\n\n==== Image File Input (OCR) ====\n\n{image_text}"
        elif file_extension == 'json':
            with open(user_input_file, 'r') as file:
                json_data = json.load(file)
                json_contents = json.dumps(json_data, indent=4)
            combined_message += f"\n\n==== JSON File Input ====\n\n{json_contents}"
        elif file_extension == 'csv':
            with open(user_input_file, mode='r', newline='', encoding='utf-8') as file:
                reader = csv.reader(file)
                csv_contents = ' '.join([','.join(row) for row in reader])
            combined_message += f"\n\n==== CSV File Input ====\n\n{csv_contents}"
        else:
            raise ValueError("Unsupported file extension")
    elif user_input_file and not os.path.exists(user_input_file):
        raise FileNotFoundError("The specified input file does not exist or is not accessible")

    # Send request to the Hugging Face API
    headers = {"Authorization": f"Bearer {api_key}"}
    response = requests.post(API_URL, headers=headers, json={"inputs": combined_message})
    if response.status_code == 200:
        response_data = response.json()
        # Save response data to JSON file
        with open(output_json_file, 'w') as jsonfile:
            json.dump(response_data, jsonfile, indent=4)
        # Extract text part of response and save to text file
        if 'generated_text' in response_data:
            with open(output_text_file, 'w') as textfile:
                textfile.write(response_data['generated_text'])
        return response_data
    else:
        return {'error': 'Failed to get a valid response', 'status_code': response.status_code}


def chatgpt_chat(model:str='gpt-3.5-turbo',
                  user_input_file:str='dataspace/user_text_input.txt',
                  output_text_file:str='dataspace/gpt_output_text.txt',
                  output_json_file:str='dataspace/gpt_output_full.json',
                  temperature:int=1, 
                  max_tokens:int=256, 
                  top_p:int=1, 
                  frequency_penalty:int=0, 
                  presence_penalty:int=0,
                  api_key:str='sk-CiO5GzpXbxZQsMuKEQEkT3BlbkFJz4LS3FuI3f5NqmF1BXO', 
                  user_message:str='Summarize:',):

    '''
    This function interfaces with OpenAI's GPT model to process a text input and obtain a generated response.
    It reads an additional input from a file, merges it with the user's direct input, and sends the combined
    content to the API. The response is then saved to both text and JSON files.

    Parameters:
        api_key (str): Your OpenAI API key.
        model (str): Identifier for the model version to use.
        user_message (str): Direct user input as a string.
        user_input_file (str): Path to a file containing additional text to send to the API.
        output_file (str): Path to save the plain text response from the API.
        output_json_file (str): Path to save the full JSON response from the API.
        temperature (float): Controls randomness in response generation. Higher is more random.
        max_tokens (int): Maximum length of the response.
        top_p (float): Controls diversity via nucleus sampling: 0.1 means only top 10% probabilities considered.
        frequency_penalty (float): Decreases likelihood of repeating words.
        presence_penalty (float): Decreases likelihood of repeating topics.

    Returns:
        str: The text response from the API.
    '''

    # Initialize variables to store responses
    ans = None
    ans_dict = {}

    combined_message = user_message  # Start with the user's direct message

    # Determine the file type and read content accordingly
    if os.path.exists(user_input_file):
        file_extension = user_input_file.split('.')[-1].lower()
        if file_extension == 'txt':
            with open(user_input_file, 'r') as file:
                file_contents = file.read().strip()
            combined_message += f"\n\n==== Text File Input ====\n\n{file_contents}"
        elif file_extension == 'pdf':
            with open(user_input_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pdf_contents = ' '.join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            combined_message += f"\n\n==== PDF File Input ====\n\n{pdf_contents}"
        elif file_extension == 'docx':
            doc = Document(user_input_file)
            docx_contents = ' '.join([para.text for para in doc.paragraphs])
            combined_message += f"\n\n==== DOCX File Input ====\n\n{docx_contents}"
        elif file_extension == 'xlsx':
            workbook = load_workbook(filename=user_input_file)
            sheet = workbook.active
            xlsx_contents = ' '.join([str(cell.value) for row in sheet for cell in row if cell.value is not None])
            combined_message += f"\n\n==== XLSX File Input ====\n\n{xlsx_contents}"
        elif file_extension in ['png', 'jpg', 'jpeg']:
            img = Image.open(user_input_file)
            image_text = pytesseract.image_to_string(img)
            combined_message += f"\n\n==== Image File Input (OCR) ====\n\n{image_text}"
        elif file_extension == 'json':
            with open(user_input_file, 'r') as file:
                json_data = json.load(file)
                json_contents = json.dumps(json_data, indent=4)
            combined_message += f"\n\n==== JSON File Input ====\n\n{json_contents}"
        elif file_extension == 'csv':
            with open(user_input_file, mode='r', newline='', encoding='utf-8') as file:
                reader = csv.reader(file)
                csv_contents = ' '.join([','.join(row) for row in reader])
            combined_message += f"\n\n==== CSV File Input ====\n\n{csv_contents}"


    try:
        # Set the API key (consider using environment variables for security)
        client = OpenAI(api_key=api_key)
        # Create a chat completion request with the specified parameters
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": combined_message}],
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=top_p,
            frequency_penalty=frequency_penalty,
            presence_penalty=presence_penalty,
        )

    # response = client.chat.completions.create(
    #     model="gpt-4o",
    #     response_format={ "type": "json_object" },
    #     messages=[
    #         {"role": "system", "content": system_message},
    #         {"role": "user", "content": user_message}
    #     ],
    #     temperature=1,
    #     max_tokens=1280,
    #     top_p=1,
    #     frequency_penalty=0,
    #     presence_penalty=0,
    #     )
    # response_text = response.choices[0].message.content.strip() # response['choices'][0]['message']['content'].strip()
    # return response_text
        # Extract and process the response
        ans_dict = response.to_dict()
        if 'choices' in ans_dict and len(ans_dict['choices']) > 0:
            if 'message' in ans_dict['choices'][0]:
                ans = ans_dict['choices'][0]['message']['content']

        # Save the text response and the full JSON response
        if ans:
            with open(output_text_file, 'w') as f:
                f.write(ans)
        with open(output_json_file, 'w') as json_file:
            json.dump(ans_dict, json_file, indent=4)

    except Exception as e:
        print(f"An unexpected error occurred: {e}")

    return ans

import pandas as pd
import plotly.graph_objects as go
from typing import List, Optional

def __plotly_chart(input_csv: str='dataspace/dataset.csv', 
                 chart_type: str='line',
                 title: str='Chart Title',
                 x_axis: str='Date',
                 y_axes: List[str] = ['Close'],  # Can be a list or a single string
                 y_secondary: Optional[str] = '',  # Optional secondary Y-axis
                 output_html: str='dataspace/plotly_viz.html'):
    """
    Generates a chart using Plotly based on the specified chart type with an optional secondary Y-axis and saves it as an HTML file.

    Parameters:
        input_csv (str): Path to the CSV file containing the data.
        chart_type (str): Type of chart to generate ('line', 'bar', 'scatter').
        x_axis (str): Column name to be used as the x-axis.
        y_axes (list or str): Column name(s) to be used as the primary y-axis. Can be a list for multiple Y-values.
        y_secondary (str, optional): Column name to be used as the secondary y-axis.
        output_html (str): Path where the HTML file will be saved.

    Returns:
        None: The function saves the chart directly to an HTML file and also returns the figure.
    """
    data = pd.read_csv(input_csv)

    # Initialize a Plotly figure
    fig = go.Figure()

    # Process primary y-axes
    if isinstance(y_axes, list):
        for y_axis in y_axes:
            if chart_type == 'line':
                fig.add_trace(go.Scatter(x=data[x_axis], y=data[y_axis], name=y_axis, yaxis='y'))
            elif chart_type == 'bar':
                fig.add_trace(go.Bar(x=data[x_axis], y=data[y_axis], name=y_axis, yaxis='y'))
            elif chart_type == 'scatter':
                fig.add_trace(go.Scatter(x=data[x_axis], y=data[y_axis], mode='markers', name=y_axis, yaxis='y'))
    else:
        if chart_type == 'line':
            fig.add_trace(go.Scatter(x=data[x_axis], y=data[y_axes], name=y_axes, yaxis='y'))
        elif chart_type == 'bar':
            fig.add_trace(go.Bar(x=data[x_axis], y=data[y_axes], name=y_axes, yaxis='y'))
        elif chart_type == 'scatter':
            fig.add_trace(go.Scatter(x=data[x_axis], y=data[y_axes], mode='markers', name=y_axes, yaxis='y'))

    # Process secondary y-axis if specified
    if y_secondary==None or y_secondary!='':
        fig.add_trace(go.Scatter(x=data[x_axis], y=data[y_secondary], name=y_secondary, yaxis='y2', marker=dict(color='red')))
        # Create a secondary y-axis configuration
        fig.update_layout(
            yaxis2=dict(
                title=y_secondary,
                overlaying='y',
                side='right'
            )
        )

    # Update layout
    fig.update_layout(
        title=title,
        xaxis_title=x_axis,
        yaxis_title=','.join(y_axes) if isinstance(y_axes, list) else y_axes
    )

    # Save the figure as an HTML file and return the figure
    fig.write_html(output_html)
    return fig
