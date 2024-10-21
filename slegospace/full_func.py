from sklearn.model_selection import train_test_split
import pandas as pd
from typing import Union
import joblib
# Import necessary libraries

def df_keep_rows_by_index(input_csv_file: str = 'dataspace/dataset.csv',
                       start_index: int = 0,
                       end_index: int = 100,
                       output_csv_file: str = 'dataspace/dataset_selected_rows.csv'):
    """
    Keeps rows in a pandas DataFrame based on specified index range.

    Parameters:
    input_csv_file (str): The file path to read the CSV file. Default is 'dataspace/dataset.csv'.
    start_index (int): The start index of the row range to keep.
    end_index (int): The end index of the row range to keep. This index is inclusive.
    output_csv_file (str): The file path to save the modified CSV file. Default is 'dataspace/selected_rows_dataset.csv'.

    Returns:
    pd.DataFrame: The DataFrame containing only the rows within the specified index range.
    """
    # Read the CSV file into a DataFrame
    df = pd.read_csv(input_csv_file)
    
    # Validate index range
    if start_index < 0 or end_index >= len(df):
        raise ValueError("Start or end index is out of the DataFrame's range.")
    
    # Keep only the rows within the specified index range
    df = df.iloc[start_index:end_index + 1]  # +1 because the end index is inclusive
    
    # Save the modified DataFrame to a CSV file
    df.to_csv(output_csv_file, index=False)
    
    return df


def df_keep_columns(input_csv_file: str = 'dataspace/dataset.csv',
                          columns_to_keep: list = ['ImportantColumn1', 'ImportantColumn2'],
                          output_csv_file: str = 'dataspace/dataset_filtered.csv'):
    """
    Keeps specified columns in a pandas DataFrame and optionally filters rows based on a condition.

    Parameters:
    input_csv_file (str): The file path to read the CSV file. Default is 'dataspace/dataset.csv'.
    columns_to_keep (list): List of column names to be retained in the DataFrame.
                            Default is ['ImportantColumn1', 'ImportantColumn2'].
    output_csv_file (str): The file path to save the modified CSV file. Default is 'dataspace/filtered_dataset.csv'.

    Returns:
    pd.DataFrame: The DataFrame with only the specified columns and optionally filtered rows.
    """
    # Read the CSV file into a DataFrame
    df = pd.read_csv(input_csv_file)
    
    # Check if all specified columns to keep are present in the DataFrame
    missing_columns = [col for col in columns_to_keep if col not in df.columns]
    if missing_columns:
        raise ValueError(f"Columns not found in DataFrame: {missing_columns}")
    
    # Keep only the specified columns
    df = df[columns_to_keep]
    
    # Save the modified DataFrame to a CSV file
    df.to_csv(output_csv_file, index=False)
    
    return df



def df_delete_columns(input_csv_file: str = 'dataspace/dataset.csv',
                   columns_to_delete: list = ['UnwantedColumn'],
                   output_csv_file: str = 'dataspace/dataset_cleaned.csv'):
    """
    Deletes specified columns from a pandas DataFrame.

    Parameters:
    input_csv_file (str): The file path to read the CSV file. Default is 'dataspace/dataset.csv'.
    columns_to_delete (list): List of column names to be deleted from the DataFrame.
                              Default is ['UnwantedColumn'].
    output_csv_file (str): The file path to save the modified CSV file. Default is 'dataspace/cleaned_dataset.csv'.

    Returns:
    pd.DataFrame: The DataFrame with the specified columns removed.
    """
    # Read the CSV file into a DataFrame
    df = pd.read_csv(input_csv_file)
    
    # Check if all specified columns to delete are present in the DataFrame
    missing_columns = [col for col in columns_to_delete if col not in df.columns]
    if missing_columns:
        raise ValueError(f"Columns not found in DataFrame: {missing_columns}")
    
    # Delete the specified columns
    df.drop(columns=columns_to_delete, inplace=True)
    
    # Save the modified DataFrame to a CSV file
    df.to_csv(output_csv_file, index=False)
    
    return df

def df_rename_columns(input_csv_file: str = 'dataspace/dataset.csv',
                      rename_dict: dict = {'Close': 'close', 'Open': 'open'},
                      output_csv_file: str = 'dataspace/dataset_renamed.csv'):
    """
    Renames multiple columns in a pandas DataFrame based on a dictionary mapping from old names to new names.

    Parameters:
    input_csv_file (str): The file path to read the CSV file. Default is 'dataspace/dataset.csv'.
    rename_dict (dict): Dictionary where keys are old column names and values are new column names.
                        Default is {'Close': 'close', 'Open': 'open'}.
    output_csv_file (str): The file path to save the modified CSV file. Default is 'dataspace/modified_dataset.csv'.

    Returns:
    pd.DataFrame: The DataFrame with renamed columns.
    """
    # Read the CSV file into a DataFrame
    df = pd.read_csv(input_csv_file)
    
    # Check if all keys in the dictionary are present in the DataFrame columns
    missing_columns = [col for col in rename_dict if col not in df.columns]
    if missing_columns:
        raise ValueError(f"Columns not found in DataFrame: {missing_columns}")
    
    # Rename the columns
    df.rename(columns=rename_dict, inplace=True)
    
    # Save the modified DataFrame to a CSV file
    df.to_csv(output_csv_file, index=False)
    
    return df

def merge_csv_by_index(input_csv_file1: str = 'dataspace/dataset1.csv',
                       input_csv_file2: str = 'dataspace/dataset2.csv',
                       output_csv_file: str = 'dataspace/merged_dataset.csv',
                       rename_dict1: dict = None,
                       rename_dict2: dict = None,
                       how: str = 'outer') -> pd.DataFrame:
    """
    Merges two CSV files horizontally based on their index after optionally renaming columns in each.

    Parameters:
    input_csv_file1 (str): Path to the first CSV file. Default path is provided.
    input_csv_file2 (str): Path to the second CSV file. Default path is provided.
    output_csv_file (str): Path where the merged CSV file will be saved. Default path is provided.
    rename_dict1 (dict): Dictionary to rename columns of the first CSV file.
    rename_dict2 (dict): Dictionary to rename columns of the second CSV file.
    how (str): Type of merge to perform ('outer', 'inner', 'left', 'right'). Default is 'outer'.

    Returns:
    pd.DataFrame: A dataframe containing the merged data.
    """
    # Read the CSV files into DataFrames
    df1 = pd.read_csv(input_csv_file1)
    df2 = pd.read_csv(input_csv_file2)

    # Rename columns if dictionaries are provided
    if rename_dict1:
        df1.rename(columns=rename_dict1, inplace=True)
    if rename_dict2:
        df2.rename(columns=rename_dict2, inplace=True)

    # Merge the DataFrames based on the index
    merged_df = pd.merge(df1, df2, left_index=True, right_index=True, how=how)
    # Drop any 'Unnamed' columns that may exist
    merged_df = merged_df.loc[:, ~merged_df.columns.str.contains('^Unnamed')]

    # Save the merged DataFrame to a CSV file
    merged_df.to_csv(output_csv_file, index=False)

    return merged_df




def prepare_dataset_tabular_ml(input_file_path: str = 'dataspace/AirQuality.csv', 
                               index_col: Union[int, bool] = 0,
                               target_column_name: str = 'NO2(GT)', 
                               drop_columns: list = ['Date','NO2(GT)','Time'],
                               output_file_path:str = 'dataspace/prepared_dataset_tabular_ml.csv'):
    """
    Prepares a dataset for tabular machine learning by loading, cleaning, and optionally modifying it.
    """
    
    # Load the dataset
    df = pd.read_csv(input_file_path, index_col=index_col)
    
    # Drop any 'Unnamed' columns that may exist
    df = df.loc[:, ~df.columns.str.contains('^Unnamed')]
    
    # Set the target column if it exists
    if target_column_name in df.columns:
        df['target'] = df[target_column_name].values
    
    # Drop specified columns if any
    if drop_columns is not None:
        df = df.drop(columns=drop_columns, errors='ignore')
    
    # Save the cleaned dataset
    output_file_path = output_file_path
    df.to_csv(output_file_path)
    
    return df



def split_dataset_4ml(input_data_file:str = 'dataspace/prepared_dataset_tabular_ml.csv', 
                           index_col: Union[int, bool] = 0,
                           train_size:int =0.6, 
                           val_size:int =0.2, 
                           test_size:int =0.2,
                           output_train_file:str ='dataspace/train.csv', 
                           output_val_file:str ='dataspace/val.csv', 
                           output_test_file:str ='dataspace/test.csv'):
    """
    Split a dataset into training, validation, and test sets and save them as CSV files.
    
    Parameters:
    - data (DataFrame): The dataset to split.
    - train_size (float): The proportion of the dataset to include in the train split.
    - val_size (float): The proportion of the dataset to include in the validation split.
    - test_size (float): The proportion of the dataset to include in the test split.
    - train_path (str): File path to save the training set CSV.
    - val_path (str): File path to save the validation set CSV.
    - test_path (str): File path to save the test set CSV.
    
    Returns:
    - train_data (DataFrame): The training set.
    - val_data (DataFrame): The validation set.
    - test_data (DataFrame): The test set.
    """

    # Load data
    data = pd.read_csv(input_data_file,index_col=index_col)

    if not (0 < train_size < 1) or not (0 < val_size < 1) or not (0 < test_size < 1):
        raise ValueError("All size parameters must be between 0 and 1.")
        
    if train_size + val_size + test_size != 1:
        raise ValueError("The sum of train_size, val_size, and test_size must equal 1.")
        
    # First split to separate out the test set
    temp_train_data, test_data = train_test_split(data, test_size=test_size, random_state=42)
    
    # Adjust val_size proportion to account for the initial split (for the remaining data)
    adjusted_val_size = val_size / (1 - test_size)
    
    # Second split to separate out the validation set from the temporary training set
    train_data, val_data = train_test_split(temp_train_data, test_size=adjusted_val_size, random_state=42)
    
    # Save to CSV
    train_data.to_csv(output_train_file, index=False)
    val_data.to_csv(output_val_file, index=False)
    test_data.to_csv(output_test_file, index=False)
    
    return train_data, val_data, test_data


def df_shift_data_row(input_csv_file: str = 'dataspace/dataset.csv',
                      shift_columns: list = ['Close'],
                      shift_columns_name: list = ['Close_shifted'],
                      shift_columns_keep: bool = False,
                      shift_rows: int = -1,
                      keep_row: bool = False,
                      output_csv_file: str = 'dataspace/dataset_selected_rows.csv'):
    """
    Shifts the data in specified columns by the specified number of rows and optionally renames and retains original columns.

    Args:
    input_csv_file (str): Path to the input CSV file.
    shift_columns (list): List of column names whose data will be shifted.
    shift_columns_name (list): List of new names for the shifted columns.
    shift_columns_keep (bool): If True, keeps the original columns alongside the shifted ones.
    shift_rows (int): Number of rows to shift the data by. 1 is one day delay, -1 is one day ahead
    keep_row (bool): If False, rows resulting in NaN values from the shift will be dropped.
    output_csv_file (str): Path to save the output CSV file after shifting the data.
    
    Returns:
    pd.DataFrame: DataFrame with the data shifted and potentially renamed in specified columns.
    """
    # Read the CSV file into a DataFrame
    df = pd.read_csv(input_csv_file)

    # Shift the data in the specified columns and optionally rename
    for orig, new in zip(shift_columns, shift_columns_name):
        df[new] = df[orig].shift(periods=shift_rows)
        if not shift_columns_keep:
            df.drop(orig, axis=1, inplace=True)

    # If keep_row is False, drop rows with NaN values that result from shifting
    if not keep_row:
        df = df.dropna(subset=shift_columns_name)

    # Save the modified DataFrame to a new CSV file
    df.to_csv(output_csv_file, index=False)

    return df










# f-image_classification.py

import os
import numpy as np
import pandas as pd
from tensorflow.keras.models import Sequential, load_model
from tensorflow.keras.layers import Conv2D, MaxPooling2D, Flatten, Dense, Dropout
from tensorflow.keras.preprocessing.image import ImageDataGenerator, img_to_array, load_img
import matplotlib.pyplot as plt
from typing import Union  # Import Union from typing
import shutil
from tensorflow.keras.datasets import cifar10  # Ensure this import is present

# CIFAR-10 Class Names
CIFAR10_CLASSES = [
    "Airplane",
    "Automobile",
    "Bird",
    "Cat",
    "Deer",
    "Dog",
    "Frog",
    "Horse",
    "Ship",
    "Truck"
]

def fetch_cifar10_dataset(
    output_data_path: str = 'dataspace/cifar10_data.npz',
    images_output_dir: str = 'dataspace/cifar10_images',
    train_split: float = 0.8,
    random_seed: int = 42
):
    """
    Fetches the CIFAR-10 dataset, saves it as a NumPy compressed file, and exports images to a directory structure.
    """
    # Load CIFAR-10 data
    (X_train, y_train), (X_test, y_test) = cifar10.load_data()

    # Save as NumPy compressed file
    np.savez_compressed(
        output_data_path,
        X_train=X_train,
        y_train=y_train,
        X_test=X_test,
        y_test=y_test
    )
    print(f"CIFAR-10 dataset saved to {output_data_path}")

    # Create directory structure
    if os.path.exists(images_output_dir):
        shutil.rmtree(images_output_dir)  # Remove existing directory to avoid duplication
    os.makedirs(images_output_dir, exist_ok=True)

    # Function to save images
    def save_images(X, y, subset):
        for idx, (img, label) in enumerate(zip(X, y)):
            class_name = CIFAR10_CLASSES[label[0]]
            class_dir = os.path.join(images_output_dir, subset, class_name)
            os.makedirs(class_dir, exist_ok=True)
            img_path = os.path.join(class_dir, f"{subset}_{idx}.png")
            plt.imsave(img_path, img)

    # Save training images
    save_images(X_train, y_train, 'train')
    print(f"Training images saved to {os.path.join(images_output_dir, 'train')}")

    # Save testing images as validation
    save_images(X_test, y_test, 'validation')
    print(f"Validation images saved to {os.path.join(images_output_dir, 'validation')}")

    return (X_train, y_train), (X_test, y_test)

def train_cnn_model(
    images_dir: str = 'dataspace/cifar10_images',
    model_output_path: str = 'dataspace/cnn_model.h5',
    performance_output_path: str = 'dataspace/cnn_performance.txt',
    epochs: int = 20,
    batch_size: int = 64,
    img_height: int = 32,
    img_width: int = 32
):
    """
    Trains a CNN model using images from a directory structure.
    """
    # Define ImageDataGenerators for training and validation
    train_datagen = ImageDataGenerator(
        rescale=1./255,
        rotation_range=15,
        width_shift_range=0.1,
        height_shift_range=0.1,
        horizontal_flip=True
    )

    validation_datagen = ImageDataGenerator(rescale=1./255)

    # Flow training images in batches
    train_generator = train_datagen.flow_from_directory(
        os.path.join(images_dir, 'train'),
        target_size=(img_height, img_width),
        batch_size=batch_size,
        class_mode='categorical',
        shuffle=True,
        seed=42
    )

    # Flow validation images in batches
    validation_generator = validation_datagen.flow_from_directory(
        os.path.join(images_dir, 'validation'),
        target_size=(img_height, img_width),
        batch_size=batch_size,
        class_mode='categorical',
        shuffle=False
    )

    # Build the CNN model
    model = Sequential([
        Conv2D(32, (3,3), activation='relu', padding='same', input_shape=(img_height, img_width, 3)),
        Conv2D(32, (3,3), activation='relu', padding='same'),
        MaxPooling2D((2,2)),
        Dropout(0.25),

        Conv2D(64, (3,3), activation='relu', padding='same'),
        Conv2D(64, (3,3), activation='relu', padding='same'),
        MaxPooling2D((2,2)),
        Dropout(0.25),

        Conv2D(128, (3,3), activation='relu', padding='same'),
        Conv2D(128, (3,3), activation='relu', padding='same'),
        MaxPooling2D((2,2)),
        Dropout(0.25),

        Flatten(),
        Dense(512, activation='relu'),
        Dropout(0.5),
        Dense(10, activation='softmax')
    ])

    model.compile(
        optimizer='adam',
        loss='categorical_crossentropy',
        metrics=['accuracy']
    )

    # Train the model
    history = model.fit(
        train_generator,
        epochs=epochs,
        validation_data=validation_generator
    )

    # Save the trained model
    model.save(model_output_path)
    print(f"Trained CNN model saved to {model_output_path}")

    # Save performance metrics
    with open(performance_output_path, 'w') as f:
        final_train_acc = history.history['accuracy'][-1]
        final_val_acc = history.history['val_accuracy'][-1]
        final_train_loss = history.history['loss'][-1]
        final_val_loss = history.history['val_loss'][-1]
        f.write(f"Final Training Accuracy: {final_train_acc:.4f}\n")
        f.write(f"Final Validation Accuracy: {final_val_acc:.4f}\n")
        f.write(f"Final Training Loss: {final_train_loss:.4f}\n")
        f.write(f"Final Validation Loss: {final_val_loss:.4f}\n")
    print(f"Performance metrics saved to {performance_output_path}")

    return history.history

def cnn_model_predict(
    model_path: str = 'dataspace/cnn_model.h5',
    input_images: Union[list, str] = 'dataspace/real_images/',
    output_data_path: str = 'dataspace/cnn_predictions.csv',
    visualize: bool = False,
    img_height: int = 32,
    img_width: int = 32
):
    """
    Uses the trained CNN model to predict classes of new images.
    """
    model = load_model(model_path)
    predictions = []

    # If input_images is a directory, get all image file paths
    if isinstance(input_images, str) and os.path.isdir(input_images):
        supported_formats = ('.png', '.jpg', '.jpeg', '.bmp', '.gif')
        image_files = [
            os.path.join(input_images, fname)
            for fname in os.listdir(input_images)
            if fname.lower().endswith(supported_formats)
        ]
    elif isinstance(input_images, list):
        image_files = input_images
    else:
        raise ValueError("input_images must be a list of file paths or a directory path.")

    if not image_files:
        print("No images found for prediction.")
        return pd.DataFrame()

    for img_path in image_files:
        try:
            img = load_img(img_path, target_size=(img_height, img_width))
            img_array = img_to_array(img) / 255.0
            img_array = np.expand_dims(img_array, axis=0)  # Shape: (1, 32, 32, 3)
            pred = model.predict(img_array)
            label_idx = np.argmax(pred, axis=1)[0]
            label_name = CIFAR10_CLASSES[label_idx]
            predictions.append({'Image_Path': img_path, 'Predicted_Label': label_name})
        except Exception as e:
            print(f"Error processing {img_path}: {e}")
            predictions.append({'Image_Path': img_path, 'Predicted_Label': 'Error'})

    df = pd.DataFrame(predictions)
    df.to_csv(output_data_path, index=False)
    print(f"Predictions saved to {output_data_path}")

    # Visualization
    if visualize:
        for _, row in df.iterrows():
            img_path = row['Image_Path']
            label = row['Predicted_Label']
            try:
                img = plt.imread(img_path)
                plt.imshow(img)
                plt.title(f"Predicted: {label}")
                plt.axis('off')
                plt.show()
            except Exception as e:
                print(f"Error displaying {img_path}: {e}")

    return df

def gather_image_paths(directory: str, supported_formats: tuple = ('.png', '.jpg', '.jpeg', '.bmp', '.gif')) -> list:
    """
    Gathers all image file paths from the specified directory.
    """
    if not os.path.isdir(directory):
        raise ValueError(f"The directory {directory} does not exist.")

    image_files = [
        os.path.join(directory, fname)
        for fname in os.listdir(directory)
        if fname.lower().endswith(supported_formats)
    ]

    if not image_files:
        print(f"No images found in directory {directory} with supported formats {supported_formats}.")

    return image_files

import pandas as pd
import vectorbt as vbt
from typing import Union
import quantstats as qs
import yfinance as yf

def moving_avg_cross_signal(input_file_path: str = 'dataspace/dataset.csv', 
                            column: str = 'Close',
                            index_col: Union[int, bool] = 0,
                            short_ma_window: int = 10,
                            long_ma_window: int = 50,
                            output_file_path:str ='dataspace/moving_avg_cross_signal.csv'

                            ):
    """Calculate moving average cross signals from a CSV file and return as a DataFrame.
    
    Args:
        input_file_path (str): Path to the CSV file containing the stock data.
        column (str): The name of the column to use for calculating moving averages.
    
    Returns:
        DataFrame: A DataFrame containing the entry and exit signals.
    """
    # Load the dataset
    df = pd.read_csv(input_file_path, index_col=index_col)
    
    # Drop any 'Unnamed' columns that may exist
    df = df.loc[:, ~df.columns.str.contains('^Unnamed')]

    # Calculate short-term and long-term moving averages
    short_ma = vbt.MA.run(df[column], window= short_ma_window)
    long_ma = vbt.MA.run(df[column], window= long_ma_window  )

    # Generate entry and exit signals
    entries = short_ma.ma_crossed_above(long_ma)
    exits = short_ma.ma_crossed_below(long_ma)

    # Convert boolean arrays to DataFrame for better handling and visualization
    signals = pd.DataFrame({
        'Entries': entries,
        'Exits': exits
    })
    #combine signal and df
    signals = pd.concat([df, signals], axis=1)

    # save csv
    signals.to_csv(output_file_path)
    return signals


def vbt_sginal_backtest(input_signal_file:str='dataspace/moving_avg_cross_signal.csv',
                             price_col:str='Close',
                             entries_col:str='Entries',
                             exits_col:str='Exits',
                             freq:str='D',
                             output_stats_file:str='dataspace/backtest_stats.csv',
                             output_return_file:str='dataspace/backtest_returns.csv'):
    """
    Create and evaluate a trading portfolio based on moving average crossover signals.
    
    This function uses vectorbt to create a portfolio from entry and exit signals
    based on moving average crossovers. It calculates the portfolio statistics and
    returns, and saves them to CSV files.

    Parameters:
        input_signal_file (str): Path to the CSV file containing price data and signals.
        price_col (str): Column name for the price data in the CSV file.
        entries_col (str): Column name for entry signals in the CSV file.
        exits_col (str): Column name for exit signals in the CSV file.
        freq (str): Frequency of the data, used for modeling in vectorbt.
        output_stats_file (str): Path where portfolio statistics will be saved.
        output_return_file (str): Path where portfolio returns will be saved.

    Returns:
        pd.DataFrame: A DataFrame containing the statistics of the portfolio.
    """

    # Load data from CSV
    data = pd.read_csv(input_signal_file)

    # Extract price and signals from the data
    price = data[price_col]
    entries = data[entries_col].astype(bool)
    exits = data[exits_col].astype(bool)

    # Create a portfolio from the signals
    portfolio = vbt.Portfolio.from_signals(price, entries, exits, freq=freq)
    
    # Save portfolio statistics to a CSV file
    stats = portfolio.stats()
    stats.to_csv(output_stats_file)

    # Compute returns and save them to a CSV file
    ret = portfolio.returns()
    ret_df = ret.to_frame(name='returns')  # Convert Series to DataFrame and name the column 'return'
    
    # Combine the original data with the returns for comprehensive output
    result = pd.concat([data, ret_df], axis=1)
    result.to_csv(output_return_file)

    # Return the statistics DataFrame
    return stats,portfolio


def backtest_viz_with_quantstats(input_file: str = 'dataspace/backtest_returns.csv', 
                                 output_file: str = 'dataspace/quantstats_results.html',
                                 benchmark_file_path:str = 'None',
                                 benchmark_col:str = 'None',
                                 return_col: str = 'returns', 
                                 time_col: str = 'Date',
                                 periods_per_year: int = 252, 
                                 compounded: str = 'True', 
                                 rf: float = 0.02,
                                 mode:str='full',
                                 title: str = 'Backtest Report Comparing Against SPY Benchmark'):
    """
    Perform a backtest visualization using QuantStats on prepared return data, optionally comparing
    to a benchmark. Defaults to SPY if no benchmark provided.

    Parameters:
        input_file (str): Path to the CSV file containing the returns data.
        output_file (str): Path to save the HTML report of the backtest results.
        return_col (str): Name of the column containing return data in the input file.
        benchmark_file_path (str, optional): Path to the CSV file containing the benchmark data.
        benchmark_col (str, optional): Column name for the benchmark data.
        periods_per_year (int): Number of periods per year for annualization (default 252 for daily data).
        compounded (bool): Whether the returns are to be treated as compounded (default True).
        rf (float): Risk-free rate used for calculating certain metrics like the Sharpe Ratio.
        time_col (str): Column containing datetime information in the input file.
        title (str): Title for the HTML report.

    Returns:
        str: Confirmation message indicating that the backtest report has been generated.
    """
    # Load historical price data from a CSV file
    data = pd.read_csv(input_file)
    data[time_col] = pd.to_datetime(data[time_col], utc=True)
    data.set_index(time_col, inplace=True)
    data.index = data.index.tz_localize(None)  # Assuming 'data' is your main DataFrame


    # Fetch SPY data from Yahoo Finance as default benchmark
    if benchmark_file_path == 'None' or benchmark_col == 'None':
        spy = yf.download('SPY', start=data.index.min(), end=data.index.max())
        benchmark = spy['Adj Close'].pct_change().dropna()
    else:
        benchmark_data = pd.read_csv(benchmark_file_path)
        benchmark_data[time_col] = pd.to_datetime(benchmark_data[time_col], utc=True)
        benchmark_data.set_index(time_col, inplace=True)
        benchmark = benchmark_data[benchmark_col]

    #rename benchamrk col
    

    # Use QuantStats to extend pandas functionality to financial series
    qs.extend_pandas()

    compounded = eval(compounded)
    # Analyze the strategy's returns and generate a report
    # Analyze the strategy's returns and generate a report
    qs.reports.html(data[return_col], 
                    benchmark=benchmark, 
                    rf=rf/periods_per_year, 
                    output=output_file, 
                    title=title, 
                    compounded=compounded, 
                    mode=mode,  # Adjust as needed
                    grayscale=False, 
                    display=False)

    return 'Backtest report generated!'

import sweetviz as sv
import pandas as pd

# from autoviz import AutoViz_Class
from typing import Union


def generate_sweetviz_report(input_csv: str='dataspace/dataset.csv', 
                             output_html: str = 'dataspace/sweetviz_report.html'):
    """
    Generates a Sweetviz report from a specified CSV file and saves it as an HTML file.

    Parameters:
    input_csv (str): Path to the input CSV dataset.
    output_html (str): Path where the HTML report will be saved, default is 'report.html'.

    Returns:
    None: The function saves the HTML report to the specified path and does not return any value.
    """
    # Load the dataset
    data = pd.read_csv(input_csv)

    # Analyze the dataset
    report = sv.analyze(data)

    # Save the report to an HTML file
    report.show_html(output_html)

    return 'Report has been generated!'



# def autoviz_plot(input_file_path: str = 'dataspace/AirQuality.csv', 
#                  target_variable: Union[str, None] = 'CO(GT)', 
#                  custom_plot_dir: str = 'dataspace',
#                  max_rows_analyzed: int = 150000,
#                  max_cols_analyzed: int = 30,
#                  lowess: bool = False,
#                  header: int = 0,
#                  verbose: int = 2,
#                  sep: str = ''):
#     """
#     Generates visualizations for the dataset using AutoViz.

#     Parameters:
#     input_file_path (str): Path to the input CSV dataset.
#     target_variable (Union[str, None]): Target variable for analysis. If None, no specific target.
#     custom_plot_dir (str): Directory where plots will be saved.
#     max_rows_analyzed (int): Maximum number of rows to analyze.
#     max_cols_analyzed (int): Maximum number of columns to analyze.
#     lowess (bool): Whether to use locally weighted scatterplot smoothing.
#     header (int): Row number to use as the column names.
#     verbose (int): Verbosity level.
#     sep (str): Separator used in the CSV file.

#     Returns:
#     str: Message indicating the completion of the visualization process.
#     """
#     AV = AutoViz_Class()

#     # Perform the AutoViz analysis and generate the plots
#     dft = AV.AutoViz(
#         filename=input_file_path,
#         sep=sep,
#         depVar=target_variable,
#         dfte=None,
#         header=header,
#         verbose=verbose,
#         lowess=lowess,
#         chart_format="html",
#         max_rows_analyzed=max_rows_analyzed,
#         max_cols_analyzed=max_cols_analyzed,
#         save_plot_dir=custom_plot_dir)

#     return "Visualizations have been generated and saved!"

import arxiv
import csv
from typing import List, Dict
import requests
import pandas as pd

import pandas as pd
import requests
import os

def download_papers_from_arxiv_csv(filename: str = "dataspace/latest_papers.csv", 
                                   download_folder: str = "dataspace/papers/",
                                   url_col: str = "entry_id",
                                   title_col: str = "title"):
    """
    Download papers from arXiv based on a CSV file.

    Parameters:
        filename (str): Path to the CSV file with arXiv paper details.
        download_folder (str): Folder to save downloaded papers.
        url_col (str): Column name in CSV that contains the arXiv URL.
        title_col (str): Column name in CSV that contains the paper title.
    """
    # Ensure the directory exists
    os.makedirs(download_folder, exist_ok=True)
    
    # Read the CSV file using Pandas
    df = pd.read_csv(filename)
    
    # Iterate over each row in the DataFrame
    for index, row in df.iterrows():
        arxiv_url = row[url_col]
        title = row[title_col].replace('/', '_').replace(':', '_')  # Clean title for filename
        arxiv_id = arxiv_url.split('/abs/')[-1].split('v')[0]  # Extract arXiv ID and remove version

        # Format the download URL and filename
        download_url = f"https://arxiv.org/pdf/{arxiv_id}.pdf"
        file_path = os.path.join(download_folder, f"{title}.pdf")

        # Download the paper
        response = requests.get(download_url)
        if response.status_code == 200:
            with open(file_path, 'wb') as f:
                f.write(response.content)
            print(f"Downloaded: {file_path}")
        else:
            print(f"Failed to download {title} with ID {arxiv_id}")

    return 'Download finished, please check your folder!'


def search_arxiv_papers(search_query: str = 'machine learning', 
                        filename: str = "dataspace/latest_papers.csv", 
                        max_results: int = 5,
                        sort_by: str = "submitted",
                        sort_order: str = "descending"):
    """
    Searches for papers on arXiv, saves the results to a CSV file, and allows sorting of the results.

    Args:
    search_query (str): The query term to search for on arXiv.
    filename (str): Path to save the CSV file containing the search results.
    max_results (int): Maximum number of results to fetch and save.
    sort_by (str): Criterion to sort the search results by ("relevance", "lastUpdatedDate", "submitted").
    sort_order (str): Order to sort the search results ("ascending", "descending").

    Returns:
    DataFrame: A pandas DataFrame containing details of the fetched papers.
    """

    # Map user-friendly sorting terms to arXiv API's SortCriterion and SortOrder
    sort_options = {
        "relevance": arxiv.SortCriterion.Relevance,
        "lastUpdatedDate": arxiv.SortCriterion.LastUpdatedDate,
        "submitted": arxiv.SortCriterion.SubmittedDate
    }
    order_options = {
        "ascending": arxiv.SortOrder.Ascending,
        "descending": arxiv.SortOrder.Descending
    }

    # Fetch the results using the arXiv API
    search = arxiv.Search(
        query=search_query,
        max_results=max_results,
        sort_by=sort_options.get(sort_by, arxiv.SortCriterion.SubmittedDate),
        sort_order=order_options.get(sort_order, arxiv.SortOrder.Descending)
    )
    results = list(search.results())

    # Convert results to DataFrame
    data = []
    for result in results:
        entry = {
            "entry_id": result.entry_id,
            "updated": result.updated.isoformat(),
            "published": result.published.isoformat(),
            "title": result.title,
            "authors": ', '.join([author.name for author in result.authors]),
            "summary": result.summary.replace('\n', ' '),
            "comment": result.comment,
            "journal_ref": result.journal_ref,
            "doi": result.doi,
            "primary_category": result.primary_category,
            "categories": ', '.join(result.categories)
        }
        data.append(entry)
    df = pd.DataFrame(data)

    # Save the DataFrame to CSV
    df.to_csv(filename, index=False)

    return df
# covid19_pipeline.py

import pandas as pd
import pickle
from prophet import Prophet
import plotly.graph_objects as go
import os

def fetch_covid19_data(
    url: str = "https://covid.ourworldindata.org/data/owid-covid-data.csv",
    output_file: str = "dataspace/covid19_data.csv",
    country: str = "United States"
) -> pd.DataFrame:
    """
    Fetches COVID-19 data for a specific country and saves it to a CSV file.

    Parameters:
        url (str): URL to fetch the COVID-19 data CSV.
        output_file (str): Path to save the filtered COVID-19 data.
        country (str): Country name to filter the data.

    Returns:
        pd.DataFrame: DataFrame containing COVID-19 data for the specified country.
    """
    df = pd.read_csv(url)
    df_country = df[df['location'] == country]
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    df_country.to_csv(output_file, index=False)
    return df_country

def preprocess_covid19_data(
    input_file: str = "dataspace/covid19_data.csv",
    output_file: str = "dataspace/covid19_preprocessed.csv",
    date_column: str = "date",
    target_column: str = "new_cases_smoothed",
    fill_missing: bool = True
) -> pd.DataFrame:
    """
    Preprocesses COVID-19 data for time-series forecasting.

    Parameters:
        input_file (str): Path to the CSV file containing COVID-19 data.
        output_file (str): Path to save the preprocessed data.
        date_column (str): Name of the column containing dates.
        target_column (str): Name of the column to forecast.
        fill_missing (bool): Whether to fill missing values.

    Returns:
        pd.DataFrame: DataFrame ready for time-series modeling.
    """
    df = pd.read_csv(input_file)
    df = df[[date_column, target_column]]
    if fill_missing:
        df[target_column].fillna(method='ffill', inplace=True)
        df[target_column].fillna(method='bfill', inplace=True)
    df = df.dropna()
    df.columns = ['ds', 'y']  # Prophet requires columns 'ds' and 'y'
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    df.to_csv(output_file, index=False)
    return df

def build_and_train_prophet_model(
    input_file: str = "dataspace/covid19_preprocessed.csv",
    model_save_path: str = "dataspace/prophet_model.pkl"
) -> Prophet:
    """
    Builds and trains a Prophet model on the preprocessed data.

    Parameters:
        input_file (str): Path to the preprocessed CSV file.
        model_save_path (str): Path to save the trained model.

    Returns:
        Prophet: Trained Prophet model.
    """
    df = pd.read_csv(input_file)
    model = Prophet()
    model.fit(df)
    os.makedirs(os.path.dirname(model_save_path), exist_ok=True)
    # Save the model to a file
    with open(model_save_path, 'wb') as f:
        pickle.dump(model, f)
    return model

def forecast_covid19_cases(
    periods: int = 30,
    model_load_path: str = "dataspace/prophet_model.pkl",
    forecast_save_path: str = "dataspace/covid19_forecast.csv"
) -> pd.DataFrame:
    """
    Uses the trained Prophet model to forecast future COVID-19 cases.

    Parameters:
        periods (int): Number of days to forecast into the future.
        model_load_path (str): Path to load the trained model.
        forecast_save_path (str): Path to save the forecasted data.

    Returns:
        pd.DataFrame: DataFrame containing forecasted values.
    """
    # Load the model
    with open(model_load_path, 'rb') as f:
        model = pickle.load(f)
    # Create future dataframe
    future = model.make_future_dataframe(periods=periods)
    forecast = model.predict(future)
    os.makedirs(os.path.dirname(forecast_save_path), exist_ok=True)
    forecast.to_csv(forecast_save_path, index=False)
    return forecast


import pandas as pd
import plotly.graph_objects as go
import os

def visualize_forecast(
    forecast_file: str = "dataspace/covid19_forecast.csv",
    actuals_file: str = "dataspace/covid19_preprocessed.csv",
    output_html_file: str = "dataspace/covid19_forecast.html"
) -> str:
    """
    Creates an interactive plot of the COVID-19 forecast alongside actual data.

    Parameters:
        forecast_file (str): Path to the CSV file containing forecasted data.
        actuals_file (str): Path to the CSV file containing actual data.
        output_html_file (str): Path to save the forecast plot as an HTML file.

    Returns:
        str: Confirmation message indicating the plot has been saved.
    """
    forecast = pd.read_csv(forecast_file)
    actuals = pd.read_csv(actuals_file)
    
    fig = go.Figure()
    # Add actual cases
    fig.add_trace(go.Scatter(
        x=actuals['ds'], y=actuals['y'], mode='markers', name='Actual Cases',
        marker=dict(color='blue')
    ))
    # Add forecasted cases
    fig.add_trace(go.Scatter(
        x=forecast['ds'], y=forecast['yhat'], mode='lines', name='Forecast',
        line=dict(color='red')
    ))
    # Add confidence intervals
    fig.add_trace(go.Scatter(
        x=forecast['ds'], y=forecast['yhat_upper'], mode='lines', name='Upper Confidence Interval',
        line=dict(width=0), showlegend=False
    ))
    fig.add_trace(go.Scatter(
        x=forecast['ds'], y=forecast['yhat_lower'], mode='lines', name='Lower Confidence Interval',
        fill='tonexty', fillcolor='rgba(68, 68, 68, 0.1)', line=dict(width=0), showlegend=False
    ))
    fig.update_layout(
        title='COVID-19 Cases Forecast',
        xaxis_title='Date',
        yaxis_title='Number of Cases',
        legend_title='Legend',
    )
    os.makedirs(os.path.dirname(output_html_file), exist_ok=True)
    fig.write_html(output_html_file)
    print(f"Forecast visualization saved to {output_html_file}")
    return "COVID-19 forecast plot saved."



# # Example usage
# if __name__ == "__main__":
#     # Step 1: Fetch data
#     df_covid = fetch_covid19_data(country="United States")

#     # Step 2: Preprocess data
#     df_preprocessed = preprocess_covid19_data()

#     # Step 3: Build and train the model
#     prophet_model = build_and_train_prophet_model()

#     # Step 4: Forecast future cases
#     forecast_df = forecast_covid19_cases(periods=30)

#     # Step 5: Visualize the forecast
#     message = visualize_forecast()
#     print(message)

import ast
import pandas as pd
import datetime
from yfinance import Ticker as si
from typing import Union, Callable
import plotly.graph_objects as go
from typing import Union

'''
Guidelines for building microservices (python functions):

1. Create a python function
2. Make sure the function has a docstring that explain to user how to use the microservice, and what the service does
3. Make sure the function has a return statement
4. Make sure the function has a parameter
5. Make sure the function has a default value for the parameter
6. Make sure the function has a type hint for the parameter 

'''
def import_marketdata_yahoo_csv(ticker: str = 'msft', 
                                start_date: str = (datetime.datetime.now() - datetime.timedelta(days=365)).strftime("%Y-%m-%d"),
                                end_date: str = datetime.datetime.now().strftime("%Y-%m-%d"), 
                                output_file_path: str = 'dataspace/dataset.csv' ):
    """
    Imports market data from Yahoo using the yfinance Ticker API.
    
    Parameters:
    - ticker (str): The stock ticker symbol. Default is 'msft'.
    - start_date (str): The start date for the data in 'YYYY-MM-DD' format. Default is one year ago from today.
    - end_date (str): The end date for the data in 'YYYY-MM-DD' format. Default is today.
    - output_file_path (str): The path to save the resulting CSV file. Default is 'yfinance_ohlc.csv'.
    
    Returns:
    - pd.DataFrame: A dataframe containing the imported market data.
    """

    # Fetch the market data using the yfinance Ticker
    df = si(ticker).history(start=start_date, end=end_date)
    df['Datetime'] = df.index
    df.to_csv(output_file_path)
    # Return the StringIO object
    return df

def preprocess_filling_missing_values(
    input_file_path: str = 'dataspace/dataset.csv',
    output_file_path: str = 'dataspace/dataset.csv',
    fill_strategy: Union[str, float, int, dict, Callable] = 'ffill',
):
    """
    Preprocesses a CSV dataset by filling missing values using various strategies and outputs the processed dataset to a CSV file.

    Parameters:
    - input_file_path (str): Path to the input CSV dataset. Defaults to 'dataspace/dataset.csv'.
    - output_file_path (str): Path for the output CSV dataset where processed data will be saved. Defaults to 'dataspace/dataset.csv'.
    - fill_strategy (Union[str, float, int, dict, Callable]): Defines the strategy for filling missing values. Can be a string ('ffill' for forward fill, 'bfill' for backward fill), a scalar value to fill missing entries, a dictionary to specify method or value per column, or a callable with custom logic. Defaults to 'ffill'.

    Returns:
    pandas.DataFrame: The processed DataFrame with missing values filled.

    This function reads a CSV file, fills missing values according to the specified strategy, and writes the processed DataFrame to the specified CSV file. It supports various filling strategies to accommodate different data preprocessing needs.
    """

    data = pd.read_csv(input_file_path)

    if callable(fill_strategy):
        # Apply a custom function to fill missing values
        data.apply(lambda x: x.fillna(fill_strategy(x)), axis=0)
    else:
        # Apply predefined pandas strategies or scalar values
        data.fillna(value=fill_strategy, inplace=True)

    data.to_csv(output_file_path)

    return data


def plot_chart_local(input_file_path: str = 'dataspace/dataset.csv', 
                     index_col: Union[None, int, str] = 0,
                     x_column: str = 'Date', 
                     y_column: str = 'SMA_Close', 
                     title: str = 'Data Plot', 
                     legend_title:str= 'Legend',
                     mode:str = 'lines',
                     output_html_file_path: str = 'dataspace/dataset_plot.html'):
    """
    Create a Plotly graph for data and save it as an HTML file locally.
    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset.csv'.
    :param x_column: Column to use for the X-axis. If None, the data index is used. Default is 'Date'.
    :param y_column: Column to use for the Y-axis. Default is 'SMA_Close'.
    :param title: Title of the plot. Default is 'Data Plot'.
    :param output_html_file_path: Local file path where the HTML file will be saved. Default is 'dataspace/dataset_plot.html'.
    """

    # Load data from a local file
    data = pd.read_csv(input_file_path, index_col=index_col)

    # Determine X-axis data
    x_data = data[x_column] if x_column else data.index

    # Create Plotly figure
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=x_data, y=data[y_column], mode=mode, name=y_column))

    # Update layout
    fig.update_layout(
        title=title,
        xaxis_title=x_column,
        yaxis_title=y_column,
        legend_title=legend_title
    )

    # Save the figure as an HTML file locally
    fig.write_html(output_html_file_path)

    return "Plot saved locally."

def compute_simple_moving_average_local(input_file_path: str = 'dataspace/dataset.csv', 
                                        column_name: str = 'Close', 
                                        window_size: int = 20, 
                                        output_file_path: str = 'dataspace/dataset.csv'):
    """
    Compute simple moving average for a specified column in the data and save updated data locally.

    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset.csv'.
    :param column_name: Name of the column to calculate the SMA on. Default is 'Close'.
    :param window_size: Window size for the moving average. Default is 20.
    :param output_file_path: Local file path where the updated data will be saved. Default is 'dataspace/dataset_sma.csv'.
    """
    # Load data from a local file
    
    data = pd.read_csv(input_file_path, index_col=0)

    # Calculate Simple Moving Average
    if column_name in data.columns:
        data[f'SMA_{column_name}'] = data[column_name].rolling(window=window_size).mean()
    else:
        raise ValueError(f"Column '{column_name}' not found in data")

    # Save updated data to a local file
    data.to_csv(output_file_path)

    return data




import ast
import pandas as pd
import datetime
from yfinance import Ticker as si
from typing import Union, Callable
import plotly.graph_objects as go
from typing import Union

'''
Guidelines for building microservices (python functions):

1. Create a python function
2. Make sure the function has a docstring that explain to user how to use the microservice, and what the service does
3. Make sure the function has a return statement
4. Make sure the function has a parameter
5. Make sure the function has a default value for the parameter
6. Make sure the function has a type hint for the parameter 

'''
def import_marketdata_yahoo_csv(ticker: str = 'msft', 
                                start_date: str = (datetime.datetime.now() - datetime.timedelta(days=365)).strftime("%Y-%m-%d"),
                                end_date: str = datetime.datetime.now().strftime("%Y-%m-%d"), 
                                output_file_path: str = 'dataspace/dataset.csv' ):
    """
    Imports market data from Yahoo using the yfinance Ticker API.
    
    Parameters:
    - ticker (str): The stock ticker symbol. Default is 'msft'.
    - start_date (str): The start date for the data in 'YYYY-MM-DD' format. Default is one year ago from today.
    - end_date (str): The end date for the data in 'YYYY-MM-DD' format. Default is today.
    - output_file_path (str): The path to save the resulting CSV file. Default is 'yfinance_ohlc.csv'.
    
    Returns:
    - pd.DataFrame: A dataframe containing the imported market data.
    """

    # Fetch the market data using the yfinance Ticker
    df = si(ticker).history(start=start_date, end=end_date)
    df['Datetime'] = df.index
    df.to_csv(output_file_path)
    # Return the StringIO object
    return df

def preprocess_filling_missing_values(
    input_file_path: str = 'dataspace/dataset.csv',
    output_file_path: str = 'dataspace/dataset.csv',
    fill_strategy: Union[str, float, int, dict, Callable] = 'ffill',
):
    """
    Preprocesses a CSV dataset by filling missing values using various strategies and outputs the processed dataset to a CSV file.

    Parameters:
    - input_file_path (str): Path to the input CSV dataset. Defaults to 'dataspace/dataset.csv'.
    - output_file_path (str): Path for the output CSV dataset where processed data will be saved. Defaults to 'dataspace/dataset.csv'.
    - fill_strategy (Union[str, float, int, dict, Callable]): Defines the strategy for filling missing values. Can be a string ('ffill' for forward fill, 'bfill' for backward fill), a scalar value to fill missing entries, a dictionary to specify method or value per column, or a callable with custom logic. Defaults to 'ffill'.

    Returns:
    pandas.DataFrame: The processed DataFrame with missing values filled.

    This function reads a CSV file, fills missing values according to the specified strategy, and writes the processed DataFrame to the specified CSV file. It supports various filling strategies to accommodate different data preprocessing needs.
    """

    data = pd.read_csv(input_file_path)

    if callable(fill_strategy):
        # Apply a custom function to fill missing values
        data.apply(lambda x: x.fillna(fill_strategy(x)), axis=0)
    else:
        # Apply predefined pandas strategies or scalar values
        data.fillna(value=fill_strategy, inplace=True)

    data.to_csv(output_file_path)

    return data

def compute_return(input_file_path: str = 'dataspace/dataset.csv',
                   output_file_path: str = 'dataspace/dataset_return.csv',
                   window_size: int = 20,
                   target_column_name: str = 'Close',
                   return_column_name: str = 'Return',
                   keep_rows: bool = False):
    """
    Compute the daily returns of a stock based on the closing price over a given window size.
    
    Args:
    input_file_path (str): Path to the input CSV file.
    output_file_path (str): Path to save the output CSV file.
    window_size (int): The number of days over which to calculate the percentage change.
    target_column_name (str): The name of the column from which to calculate returns.
    return_column_name (str): The name of the new column for the calculated returns.
    keep_rows (bool): If False, rows containing NaN values as a result of the calculation will be removed.

    Returns:
    pd.DataFrame: DataFrame with the returns added as a new column.
    """
    # Read the data from the input file
    data = pd.read_csv(input_file_path)
    
    # Calculate returns and assign them to the specified new column
    data[return_column_name] = data[target_column_name].pct_change(periods=window_size)
    
    # Handle NaN values based on keep_rows
    if not keep_rows:
        data = data.dropna(subset=[return_column_name])
    
    # Save the modified DataFrame to a new CSV file
    data.to_csv(output_file_path, index=False)
    
    return data
        

def plotly_chart(input_file_path: str = 'dataspace/dataset.csv', 
                     index_col: Union[None, int, str] = 0,
                     x_column: str = 'Date', 
                     y_column: str = 'SMA_Close', 
                     title: str = 'Data Plot', 
                     legend_title:str= 'Legend',
                     mode:str = 'lines',
                     output_html_file_path: str = 'dataspace/dataset_plot.html'):
    """
    Create a Plotly graph for data and save it as an HTML file locally.
    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset.csv'.
    :param x_column: Column to use for the X-axis. If None, the data index is used. Default is 'Date'.
    :param y_column: Column to use for the Y-axis. Default is 'SMA_Close'.
    :param title: Title of the plot. Default is 'Data Plot'.
    :param output_html_file_path: Local file path where the HTML file will be saved. Default is 'dataspace/dataset_plot.html'.
    """

    # Load data from a local file
    data = pd.read_csv(input_file_path, index_col=index_col)

    # Determine X-axis data
    x_data = data[x_column] if x_column else data.index

    # Create Plotly figure
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=x_data, y=data[y_column], mode=mode, name=y_column))

    # Update layout
    fig.update_layout(
        title=title,
        xaxis_title=x_column,
        yaxis_title=y_column,
        legend_title=legend_title
    )

    # Save the figure as an HTML file locally
    fig.write_html(output_html_file_path)

    return "Plot saved locally."

def compute_simple_moving_average(input_file_path: str = 'dataspace/dataset.csv', 
                                        column_name: str = 'Close', 
                                        window_size: int = 20, 
                                        output_file_path: str = 'dataspace/dataset.csv'):
    """
    Compute simple moving average for a specified column in the data and save updated data locally.

    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset.csv'.
    :param column_name: Name of the column to calculate the SMA on. Default is 'Close'.
    :param window_size: Window size for the moving average. Default is 20.
    :param output_file_path: Local file path where the updated data will be saved. Default is 'dataspace/dataset_sma.csv'.
    """
    # Load data from a local file
    
    data = pd.read_csv(input_file_path, index_col=0)

    # Calculate Simple Moving Average
    if column_name in data.columns:
        data[f'SMA_{column_name}'] = data[column_name].rolling(window=window_size).mean()
    else:
        raise ValueError(f"Column '{column_name}' not found in data")

    # Save updated data to a local file
    data.to_csv(output_file_path)

    return data




# import pandas as pd
# from autogluon.tabular import TabularPredictor
# import pickle

# def train_autogluon_tabular(input_file_path:str='dataspace/train.csv',
#                                target:str='target',
#                                train_frac:float=0.8,
#                                random_state:int=42,
#                                performance_output_path:str='dataspace/performance_autogluon.txt',
#                                model_save_path:str='dataspace/autogluon_model.pkl'):
#     """
#     Trains a classification model using AutoGluon on the specified dataset and saves the trained model.

#     Parameters:
#         input_file_path (str): Path to the CSV file containing the dataset.
#         target (str): Name of the column to predict.
#         train_frac (float): Fraction of the dataset to use for training.
#         random_state (int): Seed for the random number generator for reproducibility.
#         performance_output_path (str): Path to save the text file containing model performance metrics.
#         model_save_path (str): Path to save the trained AutoGluon model.

#     Returns:
#         Tuple[TabularPredictor, dict]: A tuple containing the trained AutoGluon TabularPredictor and 
#                                        a dictionary with performance metrics.

#     Saves the trained model using AutoGluon's built-in save mechanism and optionally as a pickle file.
#     Evaluates the model's performance and writes metrics to the specified text file.
#     """  
#     # Load the dataset from a CSV file
#     df = pd.read_csv(input_file_path)
#     df = df.loc[:, ~df.columns.str.contains('Unnamed: 0', case=False)]
#     train_data = df.sample(frac=train_frac, random_state=random_state)
#     test_data = df.drop(train_data.index)

#     # Train a classifier with AutoGluon
#     predictor = TabularPredictor(label=target).fit(train_data)
#     performance = predictor.evaluate(test_data)
#     leaderboard = predictor.leaderboard(test_data, silent=True)

#     # Write the performance metrics and leaderboard to a file
#     with open(performance_output_path, 'w') as f:
#         f.write(str(performance))
#         f.write("\n")
#         f.write(str(leaderboard))

#     # Save the trained model using AutoGluon's method
#     predictor.save(model_save_path)

#     # Optionally, save the model using pickle
#     pickle_path = model_save_path
#     with open(pickle_path, 'wb') as f:
#         pickle.dump(predictor, f)

#     return predictor, performance

# def autogluon_model_predict(pickle_path:str= 'dataspace/autogluon_model.pkl',
#                         input_data_path:str='dataspace/test.csv',
#                         output_data_path:str='dataspace/autogluon_predict.csv' ):
#     """
#     Loads a pickled AutoGluon model from the specified path.

#     Parameters:
#         pickle_path (str): Path to the pickled model file.

#     Returns:
#         TabularPredictor: The loaded AutoGluon model.

#     Note:
#         Loading models via pickle can lead to issues if there are mismatches in library versions or 
#         if the saved model includes elements that pickle cannot handle properly. It is generally 
#         recommended to use AutoGluon's native load functionality unless there is a specific need 
#         for pickle.
#     """

#     with open(pickle_path, 'rb') as f:
#         loaded_model = pickle.load(f)

#     predictions = loaded_model.predict(input_data_path)
#         # Save the predictions to a CSV file
#     predictions.to_csv(output_data_path, index=False)


#     return predictions


# environmental_anomaly_detection_pipeline.py

import pandas as pd
from sklearn.ensemble import IsolationForest
import plotly.graph_objects as go
import requests
import os

def fetch_air_quality_data(
    url: str = "https://archive.ics.uci.edu/ml/machine-learning-databases/00360/AirQualityUCI.zip",
    extract_to: str = "dataspace/",
    output_csv_file: str = "dataspace/AirQuality.csv"
):
    """
    Fetches the Air Quality dataset from the UCI repository, extracts it, and saves it as a CSV file.

    Parameters:
    - url (str): URL to download the dataset ZIP file.
    - extract_to (str): Directory to extract the ZIP contents.
    - output_csv_file (str): Path to save the extracted CSV file.

    Returns:
    - pd.DataFrame: The Air Quality dataset as a pandas DataFrame.
    """
    import zipfile
    import io

    # Ensure the dataspace directory exists
    os.makedirs(extract_to, exist_ok=True)

    # Download the dataset
    response = requests.get(url)
    if response.status_code == 200:
        with zipfile.ZipFile(io.BytesIO(response.content)) as z:
            # Extract the CSV file
            z.extract("AirQualityUCI.csv", path=extract_to)
        
        # Read the CSV file
        df = pd.read_csv(os.path.join(extract_to, "AirQualityUCI.csv"), sep=';', decimal=',')
        
        # Drop the last two columns which are empty
        df = df.iloc[:, :-2]
        
        # Replace comma with dot and convert to numeric
        df = df.replace(',', '.', regex=True)
        for col in df.columns[2:]:
            df[col] = pd.to_numeric(df[col], errors='coerce')
        
        # Save the cleaned dataset
        df.to_csv(output_csv_file, index=False)
        return df
    else:
        raise ConnectionError(f"Failed to download dataset. Status code: {response.status_code}")

def preprocess_air_quality_data(
    input_file_path: str = "dataspace/AirQuality.csv",
    output_file_path: str = "dataspace/preprocessed_AirQuality.csv",
    target_column: str = "NO2(GT)"
):
    """
    Preprocesses the Air Quality dataset by handling missing values and selecting relevant features.

    Parameters:
    - input_file_path (str): Path to the raw CSV dataset.
    - output_file_path (str): Path to save the preprocessed CSV dataset.
    - target_column (str): The column to analyze for anomalies.

    Returns:
    - pd.DataFrame: The preprocessed DataFrame.
    """
    df = pd.read_csv(input_file_path)
    
    # Convert 'Date' and 'Time' to datetime
    df['Datetime'] = pd.to_datetime(df['Date'] + ' ' + df['Time'], format='%d/%m/%Y %H.%M.%S', errors='coerce')
    
    # Drop rows with invalid datetime
    df.dropna(subset=['Datetime'], inplace=True)
    
    # Set datetime as index
    df.set_index('Datetime', inplace=True)
    
    # Select relevant columns (example: CO, NO2, SO2, etc.)
    relevant_columns = [
        "CO(GT)", "PT08.S1(CO)", "NMHC(GT)", "C6H6(GT)",
        "PT08.S2(NMHC)", "NOx(GT)", "PT08.S3(NOx)", "NO2(GT)",
        "PT08.S4(NO2)", "PT08.S5(O3)", "T", "RH", "AH"
    ]
    df = df[relevant_columns]
    
    # Handle missing values by forward filling
    df.fillna(method='ffill', inplace=True)
    
    # Save the preprocessed data
    df.to_csv(output_file_path)
    return df

def detect_anomalies_isolation_forest(
    input_file_path: str = "dataspace/preprocessed_AirQuality.csv",
    target_column: str = "NO2(GT)",
    output_file_path: str = "dataspace/air_quality_anomalies.csv"
):
    """
    Detects anomalies in the specified target column using Isolation Forest.

    Parameters:
    - input_file_path (str): Path to the preprocessed CSV dataset.
    - target_column (str): Name of the column to analyze for anomalies.
    - output_file_path (str): Path to save the results with anomaly labels.

    Returns:
    - pd.DataFrame: DataFrame with anomaly labels.
    """
    df = pd.read_csv(input_file_path, parse_dates=['Datetime'], index_col='Datetime')
    
    # Select the target column
    data = df[[target_column]].dropna()
    
    # Initialize Isolation Forest
    model = IsolationForest(contamination=0.01, random_state=42)
    data['Anomaly'] = model.fit_predict(data)
    
    # Convert prediction to boolean
    data['Anomaly'] = data['Anomaly'].apply(lambda x: True if x == -1 else False)
    
    # Merge anomaly labels back to the original dataframe
    df = df.join(data['Anomaly'], how='left')
    df['Anomaly'].fillna(False, inplace=True)
    
    # Save the results
    df.to_csv(output_file_path)
    return df

def plot_anomalies(
    input_file_path: str = "dataspace/air_quality_anomalies.csv",
    target_column: str = "NO2(GT)",
    date_column: str = "Datetime",
    anomaly_column: str = "Anomaly",
    output_html_file: str = "dataspace/air_quality_anomaly_plot.html"
):
    """
    Plots the time series data and highlights the detected anomalies.

    Parameters:
    - input_file_path (str): Path to the CSV file containing data and anomaly labels.
    - target_column (str): Name of the column to plot.
    - date_column (str): Name of the date column.
    - anomaly_column (str): Name of the anomaly label column.
    - output_html_file (str): Path to save the interactive Plotly HTML plot.

    Returns:
    - None: Saves the plot to the specified HTML file.
    """
    df = pd.read_csv(input_file_path, parse_dates=[date_column])
    
    # Create the main time series plot
    fig = go.Figure()
    
    fig.add_trace(
        go.Scatter(
            x=df[date_column],
            y=df[target_column],
            mode='lines',
            name=target_column
        )
    )
    
    # Highlight anomalies
    anomalies = df[df[anomaly_column]]
    fig.add_trace(
        go.Scatter(
            x=anomalies[date_column],
            y=anomalies[target_column],
            mode='markers',
            name='Anomalies',
            marker=dict(color='red', size=10, symbol='circle-open')
        )
    )
    
    # Update layout for better visualization
    fig.update_layout(
        title=f'Time Series Anomaly Detection for {target_column}',
        xaxis_title='Datetime',
        yaxis_title=target_column,
        legend=dict(x=0.01, y=0.99),
        hovermode='x unified'
    )
    
    # Save the plot as an HTML file
    fig.write_html(output_html_file)
    print(f"Anomaly plot saved to {output_html_file}")
from ucimlrepo import fetch_ucirepo 
def fetch_uci_dataset(uci_data_id:int=360, 
                      output_file_path:str='dataspace/AirQuality.csv'):


    # fetch dataset
    downloaded_data = fetch_ucirepo(id=uci_data_id)

    # data (as pandas dataframes)
    X = downloaded_data.data.features
    y = downloaded_data.data.targets

    # # metadata
    # print(air_quality.metadata)

    # # variable information
    # print(air_quality.variables)
    

    # Assuming `X` and `y` are both pandas DataFrames with the same index
    df = pd.concat([X, y], axis=1)
    df.to_csv(output_file_path)
    return df

import os
import ast
import panel as pn
import inspect
import json
def test_function(input_string:str='Hello!', 
          output_file_path:str='dataspace/output.txt'):
    """
    A simple function to save the provided input string to a specified text file and return the string.

    Parameters:
    - input_string (str): The string to be saved.
    - output_file_path (str): The file path where the string should be saved.

    Returns:
    - str: The same input string.
    """

    # Open the file at the specified path in write mode
    with open(output_file_path, 'w') as file:
        # Write the input string to the file
        file.write(input_string)

    # Return the input string
    return input_string

def _compute(module_name, input):
    module = __import__(module_name)

    #pipeline_dict = ast.literal_eval(input)
    pipeline_dict = json.loads(text)
    output = ""
    for function_name, parameters in pipeline_dict.items():
        function = eval(f"module.{function_name}")
        result = function(**parameters)

        output += "\n===================="+function_name+"====================\n\n"
        output += str(result)

    return output

def _create_multi_select_combobox(target_module):
    """
    Creates a multi-select combobox with all functions from the target_module.
    """
    
    # Get the module name (e.g., "func" if your module is named func.py)
    module_name = target_module.__name__
    
    # Get a list of all functions defined in target_module
    functions = [name for name, obj in inspect.getmembers(target_module, inspect.isfunction)
                 if obj.__module__ == module_name and not name.startswith('_')]

    # Create a multi-select combobox using the list of functions
    multi_combobox = pn.widgets.MultiChoice(name='Functions:', options=functions, height=150)

    return multi_combobox


# def _create_multi_select_combobox(func):
#   """
#   Creates a multi-select combobox with all functions from the func.py file.
#   """

#   # Get a list of all functions in the func.py file.
#   functions = [name for name, obj in inspect.getmembers(func)
#                 if inspect.isfunction(obj) and not name.startswith('_')]

#   # Create a multi-select combobox using the list of functions.
#   multi_combobox = pn.widgets.MultiChoice(name='Functions:', options=functions,  height=150)

#   return multi_combobox


def _extract_parameter(func):
    """
    Extracts the names and default values of the parameters of a function as a dictionary.

    Args:
        func: The function to extract parameter names and default values from.

    Returns:
        A dictionary where the keys are parameter names and the values are the default values.
    """
    signature = inspect.signature(func)
    parameters = signature.parameters

    parameter_dict = {}
    for name, param in parameters.items():
        if param.default != inspect.Parameter.empty:
            parameter_dict[name] = param.default
        else:
            parameter_dict[name] = None

    return parameter_dict


import os
import json

def __combine_json_files(directory, output_file):
    """
    Combine all JSON files in a directory into a single JSON file.

    Args:
    directory (str): The directory containing JSON files.
    output_file (str): The path to the output JSON file.
    """
    combined_data = []  # List to hold data from all JSON files

    # Loop over all files in the directory
    for filename in os.listdir(directory):
        if filename.endswith('.json'):  # Check for JSON files
            file_path = os.path.join(directory, filename)
            with open(file_path, 'r') as file:
                data = json.load(file)  # Load JSON data from file
                combined_data.append(data)  # Append data to the list

    # Write combined data to output JSON file
    with open(output_file, 'w') as file:
        json.dump(combined_data, file, indent=4)  # Use 'indent' for pretty-printing

    print("All JSON files have been combined into:", output_file)

# # Example usage:
# combine_json_files('/content/drive/MyDrive/SLEGO/slegospace/knowledgespace', '/content/drive/MyDrive/SLEGO/slegospace/knowledgespace/kownledge.json')

import requests
from bs4 import BeautifulSoup

def webscrape_to_txt(url: str = "https://au.finance.yahoo.com/", 
                            output_filename: str = "dataspace/output_webscrape.txt"):
    """
    Fetches the content from the specified URL and saves the textual content 
    into a text file, stripping out all HTML tags.

    Parameters:
    - url (str): The URL from which to fetch the content. Default is Yahoo Finance homepage.
    - output_filename (str): The path to the file where the text content will be saved.

    Returns:
    - None: Outputs a file with the extracted text content.
    """
    try:
        # Send a HTTP request to the URL
        response = requests.get(url)
        # Check if the request was successful
        if response.status_code == 200:
            # Parse the HTML content
            soup = BeautifulSoup(response.text, 'html.parser')
            # Extract text using .get_text()
            text_content = soup.get_text(separator='\n', strip=True)
            # Open a file in write mode
            with open(output_filename, 'w', encoding='utf-8') as file:
                file.write(text_content)
            print(f"Text content saved successfully to {output_filename}")
        else:
            print(f"Failed to retrieve the webpage. Status code: {response.status_code}")

        return text_content
    except Exception as e:
        print(f"An error occurred: {e}")


import json
import openai
import os
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import pytesseract
import csv
import json
import openai
import os
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import pytesseract
import csv
from openai import OpenAI

def chatbot_huggingface_api(API_URL: str = "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.2",
                            api_key: str = "your_api_key",
                            user_input_file: str = '',
                            output_text_file: str = 'dataspace/gpt_output_text.txt',
                            output_json_file: str = 'dataspace/gpt_output_full.json',
                            query: dict = {"inputs": "Can you please let us know more details about your "}):
    """
    Sends a query, optionally augmented with contents from various file types, to the specified Hugging Face API endpoint.
    
    This function supports processing inputs from text, PDF, DOCX, XLSX, image files (for OCR), JSON, and CSV files. The
    contents of the file are appended to a base query provided in the `query` parameter. If a file is specified but does
    not exist, the function will raise a FileNotFoundError. Unsupported file types will raise a ValueError.

    Parameters:
        API_URL (str): The URL of the Hugging Face API model to which the request will be sent.
        api_key (str): The API key required for authentication with the Hugging Face API.
        user_input_file (str): The path to the file whose contents are to be appended to the query. If empty, only the base query is used.
        output_text_file (str): The path where the plain text part of the response will be saved.
        output_json_file (str): The path where the full JSON response from the API will be saved.
        query (dict): A dictionary containing the base query. Expected to have at least a key 'inputs' with a starting string.

    Returns:
        dict: A dictionary containing the JSON response from the API. If there's an HTTP error, it returns a dictionary
              containing the error message and status code.

    Raises:
        FileNotFoundError: If the specified `user_input_file` does not exist.
        ValueError: If the file extension of `user_input_file` is not supported.

    Example Usage:
        result = chatbot_huggingface_api(
            api_key="your_api_key_here",
            user_input_file="path/to/input.txt",
            query={"inputs": "Please analyze the following data: "}
        )
        print(result)
    """
    # Function implementation here
    # Initialize the combined message with the query's input
    combined_message = query["inputs"]

    # Process file if specified and exists
    if user_input_file and os.path.exists(user_input_file):
        file_extension = user_input_file.split('.')[-1].lower()

        if file_extension == 'txt':
            with open(user_input_file, 'r') as file:
                file_contents = file.read().strip()
            combined_message += f"\n\n==== Text File Input ====\n\n{file_contents}"
        elif file_extension == 'pdf':
            with open(user_input_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pdf_contents = ' '.join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            combined_message += f"\n\n==== PDF File Input ====\n\n{pdf_contents}"
        elif file_extension == 'docx':
            doc = Document(user_input_file)
            docx_contents = ' '.join([para.text for para in doc.paragraphs])
            combined_message += f"\n\n==== DOCX File Input ====\n\n{docx_contents}"
        elif file_extension == 'xlsx':
            workbook = load_workbook(filename=user_input_file)
            sheet = workbook.active
            xlsx_contents = ' '.join([str(cell.value) for row in sheet for cell in row if cell.value is not None])
            combined_message += f"\n\n==== XLSX File Input ====\n\n{xlsx_contents}"
        elif file_extension in ['png', 'jpg', 'jpeg']:
            img = Image.open(user_input_file)
            image_text = pytesseract.image_to_string(img)
            combined_message += f"\n\n==== Image File Input (OCR) ====\n\n{image_text}"
        elif file_extension == 'json':
            with open(user_input_file, 'r') as file:
                json_data = json.load(file)
                json_contents = json.dumps(json_data, indent=4)
            combined_message += f"\n\n==== JSON File Input ====\n\n{json_contents}"
        elif file_extension == 'csv':
            with open(user_input_file, mode='r', newline='', encoding='utf-8') as file:
                reader = csv.reader(file)
                csv_contents = ' '.join([','.join(row) for row in reader])
            combined_message += f"\n\n==== CSV File Input ====\n\n{csv_contents}"
        else:
            raise ValueError("Unsupported file extension")
    elif user_input_file and not os.path.exists(user_input_file):
        raise FileNotFoundError("The specified input file does not exist or is not accessible")

    # Send request to the Hugging Face API
    headers = {"Authorization": f"Bearer {api_key}"}
    response = requests.post(API_URL, headers=headers, json={"inputs": combined_message})
    if response.status_code == 200:
        response_data = response.json()
        # Save response data to JSON file
        with open(output_json_file, 'w') as jsonfile:
            json.dump(response_data, jsonfile, indent=4)
        # Extract text part of response and save to text file
        if 'generated_text' in response_data:
            with open(output_text_file, 'w') as textfile:
                textfile.write(response_data['generated_text'])
        return response_data
    else:
        return {'error': 'Failed to get a valid response', 'status_code': response.status_code}


def chatgpt_chat(model:str='gpt-3.5-turbo',
                  user_input_file:str='dataspace/user_text_input.txt',
                  output_text_file:str='dataspace/gpt_output_text.txt',
                  output_json_file:str='dataspace/gpt_output_full.json',
                  temperature:int=1, 
                  max_tokens:int=256, 
                  top_p:int=1, 
                  frequency_penalty:int=0, 
                  presence_penalty:int=0,
                  api_key:str='sk-CiO5GzpXbxZQsMuKEQEkT3BlbkFJz4LS3FuI3f5NqmF1BXO', 
                  user_message:str='Summarize:',):

    '''
    This function interfaces with OpenAI's GPT model to process a text input and obtain a generated response.
    It reads an additional input from a file, merges it with the user's direct input, and sends the combined
    content to the API. The response is then saved to both text and JSON files.

    Parameters:
        api_key (str): Your OpenAI API key.
        model (str): Identifier for the model version to use.
        user_message (str): Direct user input as a string.
        user_input_file (str): Path to a file containing additional text to send to the API.
        output_file (str): Path to save the plain text response from the API.
        output_json_file (str): Path to save the full JSON response from the API.
        temperature (float): Controls randomness in response generation. Higher is more random.
        max_tokens (int): Maximum length of the response.
        top_p (float): Controls diversity via nucleus sampling: 0.1 means only top 10% probabilities considered.
        frequency_penalty (float): Decreases likelihood of repeating words.
        presence_penalty (float): Decreases likelihood of repeating topics.

    Returns:
        str: The text response from the API.
    '''

    # Initialize variables to store responses
    ans = None
    ans_dict = {}

    combined_message = user_message  # Start with the user's direct message

    # Determine the file type and read content accordingly
    if os.path.exists(user_input_file):
        file_extension = user_input_file.split('.')[-1].lower()
        if file_extension == 'txt':
            with open(user_input_file, 'r') as file:
                file_contents = file.read().strip()
            combined_message += f"\n\n==== Text File Input ====\n\n{file_contents}"
        elif file_extension == 'pdf':
            with open(user_input_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pdf_contents = ' '.join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            combined_message += f"\n\n==== PDF File Input ====\n\n{pdf_contents}"
        elif file_extension == 'docx':
            doc = Document(user_input_file)
            docx_contents = ' '.join([para.text for para in doc.paragraphs])
            combined_message += f"\n\n==== DOCX File Input ====\n\n{docx_contents}"
        elif file_extension == 'xlsx':
            workbook = load_workbook(filename=user_input_file)
            sheet = workbook.active
            xlsx_contents = ' '.join([str(cell.value) for row in sheet for cell in row if cell.value is not None])
            combined_message += f"\n\n==== XLSX File Input ====\n\n{xlsx_contents}"
        elif file_extension in ['png', 'jpg', 'jpeg']:
            img = Image.open(user_input_file)
            image_text = pytesseract.image_to_string(img)
            combined_message += f"\n\n==== Image File Input (OCR) ====\n\n{image_text}"
        elif file_extension == 'json':
            with open(user_input_file, 'r') as file:
                json_data = json.load(file)
                json_contents = json.dumps(json_data, indent=4)
            combined_message += f"\n\n==== JSON File Input ====\n\n{json_contents}"
        elif file_extension == 'csv':
            with open(user_input_file, mode='r', newline='', encoding='utf-8') as file:
                reader = csv.reader(file)
                csv_contents = ' '.join([','.join(row) for row in reader])
            combined_message += f"\n\n==== CSV File Input ====\n\n{csv_contents}"


    try:
        # Set the API key (consider using environment variables for security)
        client = OpenAI(api_key=api_key)
        # Create a chat completion request with the specified parameters
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": combined_message}],
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=top_p,
            frequency_penalty=frequency_penalty,
            presence_penalty=presence_penalty,
        )

    # response = client.chat.completions.create(
    #     model="gpt-4o",
    #     response_format={ "type": "json_object" },
    #     messages=[
    #         {"role": "system", "content": system_message},
    #         {"role": "user", "content": user_message}
    #     ],
    #     temperature=1,
    #     max_tokens=1280,
    #     top_p=1,
    #     frequency_penalty=0,
    #     presence_penalty=0,
    #     )
    # response_text = response.choices[0].message.content.strip() # response['choices'][0]['message']['content'].strip()
    # return response_text
        # Extract and process the response
        ans_dict = response.to_dict()
        if 'choices' in ans_dict and len(ans_dict['choices']) > 0:
            if 'message' in ans_dict['choices'][0]:
                ans = ans_dict['choices'][0]['message']['content']

        # Save the text response and the full JSON response
        if ans:
            with open(output_text_file, 'w') as f:
                f.write(ans)
        with open(output_json_file, 'w') as json_file:
            json.dump(ans_dict, json_file, indent=4)

    except Exception as e:
        print(f"An unexpected error occurred: {e}")

    return ans

import pandas as pd
import plotly.graph_objects as go
from typing import List, Optional

def __plotly_chart(input_csv: str='dataspace/dataset.csv', 
                 chart_type: str='line',
                 title: str='Chart Title',
                 x_axis: str='Date',
                 y_axes: List[str] = ['Close'],  # Can be a list or a single string
                 y_secondary: Optional[str] = '',  # Optional secondary Y-axis
                 output_html: str='dataspace/plotly_viz.html'):
    """
    Generates a chart using Plotly based on the specified chart type with an optional secondary Y-axis and saves it as an HTML file.

    Parameters:
        input_csv (str): Path to the CSV file containing the data.
        chart_type (str): Type of chart to generate ('line', 'bar', 'scatter').
        x_axis (str): Column name to be used as the x-axis.
        y_axes (list or str): Column name(s) to be used as the primary y-axis. Can be a list for multiple Y-values.
        y_secondary (str, optional): Column name to be used as the secondary y-axis.
        output_html (str): Path where the HTML file will be saved.

    Returns:
        None: The function saves the chart directly to an HTML file and also returns the figure.
    """
    data = pd.read_csv(input_csv)

    # Initialize a Plotly figure
    fig = go.Figure()

    # Process primary y-axes
    if isinstance(y_axes, list):
        for y_axis in y_axes:
            if chart_type == 'line':
                fig.add_trace(go.Scatter(x=data[x_axis], y=data[y_axis], name=y_axis, yaxis='y'))
            elif chart_type == 'bar':
                fig.add_trace(go.Bar(x=data[x_axis], y=data[y_axis], name=y_axis, yaxis='y'))
            elif chart_type == 'scatter':
                fig.add_trace(go.Scatter(x=data[x_axis], y=data[y_axis], mode='markers', name=y_axis, yaxis='y'))
    else:
        if chart_type == 'line':
            fig.add_trace(go.Scatter(x=data[x_axis], y=data[y_axes], name=y_axes, yaxis='y'))
        elif chart_type == 'bar':
            fig.add_trace(go.Bar(x=data[x_axis], y=data[y_axes], name=y_axes, yaxis='y'))
        elif chart_type == 'scatter':
            fig.add_trace(go.Scatter(x=data[x_axis], y=data[y_axes], mode='markers', name=y_axes, yaxis='y'))

    # Process secondary y-axis if specified
    if y_secondary==None or y_secondary!='':
        fig.add_trace(go.Scatter(x=data[x_axis], y=data[y_secondary], name=y_secondary, yaxis='y2', marker=dict(color='red')))
        # Create a secondary y-axis configuration
        fig.update_layout(
            yaxis2=dict(
                title=y_secondary,
                overlaying='y',
                side='right'
            )
        )

    # Update layout
    fig.update_layout(
        title=title,
        xaxis_title=x_axis,
        yaxis_title=','.join(y_axes) if isinstance(y_axes, list) else y_axes
    )

    # Save the figure as an HTML file and return the figure
    fig.write_html(output_html)
    return fig
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import plotly.express as px
import plotly.graph_objects as go

import vectorbt as vbt
# caclulate moviang average of stock price

def movingAvg(input_file_path: str = 'dataspace/dataset.csv', 
                                        column_name: str = 'Close', 
                                        window_size: int = 50, 
                                        output_file_path: str = 'dataspace/dataset_avg50.csv'):
    """
    Compute simple moving average for a specified column in the data and save updated data locally.

    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset.csv'.
    :param column_name: Name of the column to calculate the SMA on. Default is 'Close'.
    :param window_size: Window size for the moving average. Default is 20.
    :param output_file_path: Local file path where the updated data will be saved. Default is 'dataspace/dataset_sma.csv'.
    """
    # Load data from a local file
    
    data = pd.read_csv(input_file_path, index_col=0)

    # Calculate Simple Moving Average
    if column_name in data.columns:
        data[f'SMA_{column_name}'] = data[column_name].rolling(window=window_size).mean()
    else:
        raise ValueError(f"Column '{column_name}' not found in data")

    # Save updated data to a local file
    data.to_csv(output_file_path)

    return data

def movingAvg_50(input_file_path: str = 'dataspace/dataset.csv', 
                                        column_name: str = 'Close', 
                                        window_size: int = 50, 
                                        output_file_path: str = 'dataspace/dataset_avg50.csv'):
    """
    Compute simple moving average for a specified column in the data and save updated data locally.

    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset.csv'.
    :param column_name: Name of the column to calculate the SMA on. Default is 'Close'.
    :param window_size: Window size for the moving average. Default is 20.
    :param output_file_path: Local file path where the updated data will be saved. Default is 'dataspace/dataset_sma.csv'.
    """
    # Load data from a local file
    
    data = pd.read_csv(input_file_path, index_col=0)

    # Calculate Simple Moving Average
    if column_name in data.columns:
        data[f'SMA_{column_name}'] = data[column_name].rolling(window=window_size).mean()
    else:
        raise ValueError(f"Column '{column_name}' not found in data")

    # Save updated data to a local file
    data.to_csv(output_file_path)

    return data

def movingAvg_10(input_file_path: str = 'dataspace/dataset.csv', 
                                        column_name: str = 'Close', 
                                        window_size: int = 10, 
                                        output_file_path: str = 'dataspace/dataset_avg10.csv'):
    """
    Compute simple moving average for a specified column in the data and save updated data locally.

    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset.csv'.
    :param column_name: Name of the column to calculate the SMA on. Default is 'Close'.
    :param window_size: Window size for the moving average. Default is 20.
    :param output_file_path: Local file path where the updated data will be saved. Default is 'dataspace/dataset_sma.csv'.
    """
    # Load data from a local file
    
    data = pd.read_csv(input_file_path, index_col=0)

    # Calculate Simple Moving Average
    if column_name in data.columns:
        data[f'SMA_{column_name}'] = data[column_name].rolling(window=window_size).mean()
    else:
        raise ValueError(f"Column '{column_name}' not found in data")

    # Save updated data to a local file
    data.to_csv(output_file_path)

    return data



def calculate_vwap(input_file_path: str = 'dataspace/dataset.csv', 
                   price_column: str = 'Close', 
                   volume_column: str = 'Volume', 
                   output_file_path: str = 'dataspace/dataset_vwap.csv'):
    """
    Calculate the Volume Weighted Average Price (VWAP) and save the result to a CSV file.
    
    Parameters:
    input_file_path (str): Path to the input CSV file.
    price_column (str): Name of the column containing the price data.
    volume_column (str): Name of the column containing the volume data.
    output_file_path (str): Path to the output CSV file where the data with VWAP will be saved.
    
    Returns:
    pd.DataFrame: DataFrame with the VWAP column added.
    """
    # Load data from a local file
    data = pd.read_csv(input_file_path, index_col=0)

    # Check if the required columns are in the DataFrame
    if price_column not in data.columns or volume_column not in data.columns:
        raise ValueError(f"DataFrame must contain '{price_column}' and '{volume_column}' columns")

    # Calculate VWAP
    data['VWAP'] = (data[price_column] * data[volume_column]).cumsum() / data[volume_column].cumsum()

    # Save updated data to a local file
    data.to_csv(output_file_path)

    return data


def generate_mAvg_Crossover_trading_signals(price_file_path: str = 'dataspace/dataset.csv',
                                                               price_column: str = 'Close',
                                                               short_ma_file_path: str = 'dataspace/dataset_avg10.csv',
                                                               long_ma_file_path: str = 'dataspace/dataset_avg50.csv',
                                                               output_file_path: str = 'dataspace/dataset_signals.csv'):
    """
    Generate trading signals based on moving average crossover strategy using pre-calculated moving averages
    and save the result to a CSV file.

    Parameters:
    price_file (str): Path to the CSV file containing the price data.
    short_ma_file (str): Path to the CSV file containing the short-term moving average.
    long_ma_file (str): Path to the CSV file containing the long-term moving average.
    price_column (str): Name of the column containing the price data.
    output_file_path (str): Path to the output CSV file where the data with trading signals will be saved.

    Returns:
    pd.DataFrame: DataFrame with the trading signals added.
    """
    # Load data from the price and moving average files
    price_data = pd.read_csv(price_file_path, index_col=0)
    short_ma_data = pd.read_csv(short_ma_file_path, index_col=0)
    long_ma_data = pd.read_csv(long_ma_file_path, index_col=0)

    # Check if the price column is in the price DataFrame
    if price_column not in price_data.columns:
        raise ValueError(f"The price DataFrame must contain '{price_column}' column")

    # Merge data on the index (date)
    data = price_data[[price_column]].join(short_ma_data[['SMA_Close']].rename(columns={'SMA_Close': 'SMA_Short'}))
    data = data.join(long_ma_data[['SMA_Close']].rename(columns={'SMA_Close': 'SMA_Long'}))

    # Generate trading signals
    data['Signal'] = np.where(data['SMA_Short'] > data['SMA_Long'], 1, 0)

    # Generate trading positions
    data['Position'] = data['Signal'].diff()

    # Save updated data to a local file
    data.to_csv(output_file_path)

    return data


def calculate_typical_price(input_file_path: str = 'dataspace/dataset.csv',
                            output_file_path: str = 'dataspace/dataset_typical_price.csv'):
    """
    Calculate the Typical Price for each period and save the updated data locally.

    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset.csv'.
    :param output_file_path: Local file path where the updated data will be saved. Default is 'dataspace/dataset_typical_price.csv'.
    """
    # Load data from a local file
    data = pd.read_csv(input_file_path, index_col=0)
    
    # Calculate Typical Price
    if {'High', 'Low', 'Close'}.issubset(data.columns):
        data['Typical_Price'] = (data['High'] + data['Low'] + data['Close']) / 3
    else:
        raise ValueError("Data must contain 'High', 'Low', and 'Close' columns")

    # Save updated data to a local file
    data.to_csv(output_file_path)
    
    return data

def calculate_raw_money_flow_RMF(input_file_path: str = 'dataspace/dataset_typical_price.csv',
                             output_file_path: str = 'dataspace/dataset_raw_money_flow.csv'):
    """
    Calculate the Raw Money Flow for each period and save the updated data locally.

    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset_typical_price.csv'.
    :param output_file_path: Local file path where the updated data will be saved. Default is 'dataspace/dataset_raw_money_flow.csv'.
    """
    # Load data from a local file
    data = pd.read_csv(input_file_path, index_col=0)
    
    # Calculate Raw Money Flow
    if 'Typical_Price' in data.columns and 'Volume' in data.columns:
        data['Raw_Money_Flow'] = data['Typical_Price'] * data['Volume']
    else:
        raise ValueError("Data must contain 'Typical_Price' and 'Volume' columns")
    
    # Save updated data to a local file
    data.to_csv(output_file_path)
    
    return data

def calculate_money_flow_ratio_MFR(input_file_path: str = 'dataspace/dataset_raw_money_flow.csv',
                         period: int = 14,
                         output_file_path: str = 'dataspace/dataset_money_flow_ratio.csv'):
    """
    Calculate the Positive and Negative Money Flow for each period and save the updated data locally.

    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset_raw_money_flow.csv'.
    :param period: The period for calculating the rolling sums. Default is 14.
    :param output_file_path: Local file path where the updated data will be saved. Default is 'dataspace/dataset_money_flow.csv'.
    """
    # Load data from a local file
    data = pd.read_csv(input_file_path, index_col=0)
    
    # Calculate Positive and Negative Money Flow
    data['Positive_Flow'] = 0
    data['Negative_Flow'] = 0
    
    for i in range(1, len(data)):
        if data['Typical_Price'].iloc[i] > data['Typical_Price'].iloc[i - 1]:
            data.at[data.index[i], 'Positive_Flow'] = data['Raw_Money_Flow'].iloc[i]
        elif data['Typical_Price'].iloc[i] < data['Typical_Price'].iloc[i - 1]:
            data.at[data.index[i], 'Negative_Flow'] = data['Raw_Money_Flow'].iloc[i]
    
    data['Positive_Flow'] = data['Positive_Flow'].rolling(window=period).sum()
    data['Negative_Flow'] = data['Negative_Flow'].rolling(window=period).sum()
    
    # Save updated data to a local file
    data.to_csv(output_file_path)
    
    return data

def calculate_money_flow_index_MFI(input_file_path: str = 'dataspace/dataset_money_flow_ratio.csv',
                  output_file_path: str = 'dataspace/dataset_mfi.csv'):
    """
    Calculate the Money Flow Index (MFI) and save the updated data locally.

    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset_money_flow.csv'.
    :param period: The period for calculating the MFI. Default is 14.
    :param output_file_path: Local file path where the updated data will be saved. Default is 'dataspace/dataset_mfi.csv'.
    """
    # Load data from a local file
    data = pd.read_csv(input_file_path, index_col=0)
    
    # Calculate Money Flow Ratio and MFI
    data['Money_Flow_Ratio'] = data['Positive_Flow'] / data['Negative_Flow']
    data['MFI'] = 100 - (100 / (1 + data['Money_Flow_Ratio']))
    
    # Save updated data to a local file
    data.to_csv(output_file_path)
    
    return data


def generate_mfi_trading_signals(input_file_path: str = 'dataspace/dataset_mfi.csv',
                                 mfi_column: str = 'MFI',
                                 threshold_overbought: int = 80,
                                 threshold_oversold: int = 20,
                                 output_file_path: str = 'dataspace/dataset_mfi_signals.csv'):
    """
    Generate trading signals based on the Money Flow Index (MFI) and save the result to a CSV file.

    :param input_file_path: Local file path for the input data file. Default is 'dataspace/dataset_mfi.csv'.
    :param mfi_column: Name of the column containing the MFI data. Default is 'MFI'.
    :param threshold_overbought: Threshold for overbought condition. Default is 80.
    :param threshold_oversold: Threshold for oversold condition. Default is 20.
    :param output_file_path: Local file path where the data with trading signals will be saved. Default is 'dataspace/dataset_mfi_signals.csv'.
    :return: pd.DataFrame with the trading signals added.
    """
    # Load data from a local file
    data = pd.read_csv(input_file_path, index_col=0)
    
    # Check if the MFI column is in the DataFrame
    if mfi_column not in data.columns:
        raise ValueError(f"DataFrame must contain '{mfi_column}' column")
    
    # Generate trading signals
    data['Signal'] = 0
    data.loc[data[mfi_column] > threshold_overbought, 'Signal'] = -1  # Sell signal
    data.loc[data[mfi_column] < threshold_oversold, 'Signal'] = 1     # Buy signal

    # Generate trading positions
    data['Position'] = data['Signal'].diff()
    
    # Save updated data to a local file
    data.to_csv(output_file_path)
    
    return data


def backtest_signals(signal_path='dataspace/dataset_signals.csv', 
                     signal_col='Signal', 
                     price_path='dataspace/dataset.csv', 
                     price_col='Close', 
                     initial_cash=100000, 
                     stats_output='dataspace/pf_stats.txt', 
                     plot_output='dataspace/pf_plot.html'):
    price_data = pd.read_csv(price_path, index_col=0, parse_dates=True)
    signals = pd.read_csv(signal_path, index_col=0, parse_dates=True)
    price_data = price_data.loc[signals.index]
    entries = signals[signal_col] == 1
    exits = signals[signal_col] == 0
    pf = vbt.Portfolio.from_signals(price_data[price_col], entries, exits, init_cash=initial_cash)
    stats = pf.stats()
    with open(stats_output, 'w') as f:
        f.write(stats.to_string())
    pf.plot().write_html(plot_output)
    return pf

import pandas as pd
import numpy as np

def mix_signal_strategy(signal1_path: str = 'dataspace/dataset_signals.csv',
                        signal2_path: str = 'dataspace/dataset_mfi_signals.csv',
                        signal1_name: str = 'Signal1',
                        signal2_name: str = 'Signal2',
                        output_file_path: str = 'dataspace/dataset_mixed_signals.csv',
                        buy_strategy: str = 'Signal1:1 AND Signal2:1',
                        sell_strategy: str = 'Signal1:-1 OR Signal2:-1'):
    """
    Generate trading signals by combining any two signal strategies.
    Users can specify any combination of conditions for buy and sell signals using a simple string format.

    :param signal1_path: Path to the CSV file containing the first signal. Default is 'dataspace/dataset_signals.csv'.
    :param signal2_path: Path to the CSV file containing the second signal. Default is 'dataspace/dataset_mfi_signals.csv'.
    :param signal1_name: Name to use for the first signal in strategy strings. Default is 'Signal1'.
    :param signal2_name: Name to use for the second signal in strategy strings. Default is 'Signal2'.
    :param output_file_path: Path to save the combined signals CSV file. Default is 'dataspace/dataset_mixed_signals.csv'.
    :param buy_strategy: Strategy for buy signals. Default is 'Signal1:1 AND Signal2:1'.
    :param sell_strategy: Strategy for sell signals. Default is 'Signal1:-1 OR Signal2:-1'.
    :return: pd.DataFrame with the combined trading signals.
    """
    # Load signals
    signal1_data = pd.read_csv(signal1_path, index_col=0, parse_dates=True)
    signal2_data = pd.read_csv(signal2_path, index_col=0, parse_dates=True)

    # Ensure both DataFrames have the same index
    common_index = signal1_data.index.intersection(signal2_data.index)
    signal1_data = signal1_data.loc[common_index]
    signal2_data = signal2_data.loc[common_index]

    # Create a new DataFrame for combined signals
    combined_signals = pd.DataFrame(index=common_index)
    combined_signals[signal1_name] = signal1_data['Signal']
    combined_signals[signal2_name] = signal2_data['Signal']

    def parse_strategy(strategy_str):
        conditions = strategy_str.split()
        parsed_condition = []
        for condition in conditions:
            if ':' in condition:
                signal, value = condition.split(':')
                parsed_condition.append(f"(combined_signals['{signal}'] == {value})")
            elif condition.upper() in ['AND', 'OR']:
                parsed_condition.append(condition.lower())
            else:
                parsed_condition.append(condition)
        return ' '.join(parsed_condition)

    buy_condition = parse_strategy(buy_strategy)
    sell_condition = parse_strategy(sell_strategy)

    # Generate combined signals based on the specified strategies
    combined_signals['Signal'] = np.select(
        [combined_signals.eval(buy_condition), combined_signals.eval(sell_condition)],
        [1, -1],
        default=0  # Do nothing if neither buy nor sell conditions are met
    )

    # Generate trading positions (changes in signal)
    combined_signals['Position'] = combined_signals['Signal'].diff()

    # Save combined signals to a CSV file
    combined_signals.to_csv(output_file_path)

    return combined_signals
# AUTHOR: Bhushan Oza
import pandas as pd
import datetime                            
import os
from datetime import datetime
import numpy as np
from math import log

def populate(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_populate.csv'):
    df = pd.read_csv(input_file_path)
    df = df[['Date-Time','Type','Price','Volume','Qualifiers','#RIC', 'Bid Price', 'Ask Price', 'Bid Size', 'Ask Size']]
    df = df.drop(df[df.Type != 'Trade'].index)
    df = df.drop(df[pd.isna(df.Price)].index)
    df['PV'] = df.Price * df.Volume
    for i in range(len(df['Date-Time'])):
        df.iat[i, df.columns.get_loc('Date-Time')] = (df.iat[i, df.columns.get_loc('Date-Time')])[0:10]+' '+(df.iat[i, df.columns.get_loc('Date-Time')])[11:19]
    period = input('Enter aggregation period in minutes: ')
    print('Selected period is '+str(period)+' minutes.')  
    df['Date-Time'] = pd.to_datetime(df['Date-Time']) 
    df.to_csv(output_file_path)
    return df

def dollarVolumeTraded(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_dollarVolumeTraded.csv'):
    df = populate(input_file_path)
    newDf = pd.DataFrame() 
    df_b = df.drop(df[df.Qualifiers.str[0] != 'B'].index)
    df_s = df.drop(df[df.Qualifiers.str[0] != 'S'].index)
    summed = df.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    summed_b = df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    summed_s = df_s.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    newDf['DollarVolumeTraded'] = summed.iloc[:,2]
    cats = ['b','s']
    for i in cats:
        if i=='b':
            newDf['DollarVolumeTraded_<'+i+'>'] = summed_b.iloc[:,2]
        elif i=='s':
            newDf['DollarVolumeTraded_<'+i+'>'] = summed_s.iloc[:,2]
    newDf = newDf.fillna('Not Applicable')
    newDf.to_csv(output_file_path)
    return newDf
    
def shareVolumeTraded(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_shareVolumeTraded.csv'):
    df = populate(input_file_path)
    newDf = pd.DataFrame() 
    df_b = df.drop(df[df.Qualifiers.str[0] != 'B'].index)
    df_s = df.drop(df[df.Qualifiers.str[0] != 'S'].index)
    summed = df.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    summed_b = df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    summed_s = df_s.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    cats = ['b','s']
    newDf['ShareVolumeTraded'] = summed.iloc[:,1]
    for i in cats:
        if i=='b':
            newDf['ShareVolumeTraded_<'+i+'>'] = summed_b.iloc[:,1]
        elif i=='s':
            newDf['ShareVolumeTraded_<'+i+'>'] = summed_s.iloc[:,1]
    newDf = newDf.fillna('Not Applicable')
    newDf.to_csv(output_file_path)
    return newDf  

def vWAP(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_vWAP.csv'):
    df= populate(input_file_path)
    newDf = pd.DataFrame() 
    df_b = df.drop(df[df.Qualifiers.str[0] != 'B'].index)
    df_s = df.drop(df[df.Qualifiers.str[0] != 'S'].index)
    summed = df.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    summed_b = df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    summed_s = df_s.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    newDf['DollarVolumeTraded'] = summed.iloc[:,2]
    cats = ['b','s']
    for i in cats:
        if i=='b':
            newDf['DollarVolumeTraded_<'+i+'>'] = summed_b.iloc[:,2]
        elif i=='s':
            newDf['DollarVolumeTraded_<'+i+'>'] = summed_s.iloc[:,2]
    newDf['ShareVolumeTraded'] = summed.iloc[:,1]
    for i in cats:
        if i=='b':
            newDf['ShareVolumeTraded_<'+i+'>'] = summed_b.iloc[:,1]
        elif i=='s':
            newDf['ShareVolumeTraded_<'+i+'>'] = summed_s.iloc[:,1]
    newDf['VWAP'] = newDf.DollarVolumeTraded / newDf.ShareVolumeTraded
    for i in cats:
        if i=='b':
            newDf['VWAP_<'+i+'>'] = newDf['DollarVolumeTraded_<b>'] / newDf['ShareVolumeTraded_<b>']
        elif i=='s':
            newDf['VWAP_<'+i+'>'] = newDf['DollarVolumeTraded_<s>'] / newDf['ShareVolumeTraded_<s>']  
    newDf = newDf.fillna('Not Applicable')
    newDf.to_csv(output_file_path)
    return newDf

def arithmeticReturn(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_arithmeticReturn.csv'):
    df = populate(input_file_path)
    newDf = pd.DataFrame() 
    df_b = df.drop(df[df.Qualifiers.str[0] != 'B'].index)
    df_s = df.drop(df[df.Qualifiers.str[0] != 'S'].index)
    summed = df.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    summed_b = df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    summed_s = df_s.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    newDf['DollarVolumeTraded'] = summed.iloc[:,2]
    cats = ['b','s']
    for i in cats:
        if i=='b':
            newDf['DollarVolumeTraded_<'+i+'>'] = summed_b.iloc[:,2]
        elif i=='s':
            newDf['DollarVolumeTraded_<'+i+'>'] = summed_s.iloc[:,2]
    newDf['ShareVolumeTraded'] = summed.iloc[:,1]
    for i in cats:
        if i=='b':
            newDf['ShareVolumeTraded_<'+i+'>'] = summed_b.iloc[:,1]
        elif i=='s':
            newDf['ShareVolumeTraded_<'+i+'>'] = summed_s.iloc[:,1]
    newDf['VWAP'] = newDf.DollarVolumeTraded / newDf.ShareVolumeTraded
    for i in cats:
        if i=='b':
            newDf['VWAP_<'+i+'>'] = newDf['DollarVolumeTraded_<b>'] / newDf['ShareVolumeTraded_<b>']
        elif i=='s':
            newDf['VWAP_<'+i+'>'] = newDf['DollarVolumeTraded_<s>'] / newDf['ShareVolumeTraded_<s>']  
    newDf['AReturn'] = newDf.VWAP.pct_change()
    newDf = newDf.fillna('Not Applicable')
    newDf.to_csv(output_file_path)
    return newDf

def logReturn(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_logReturn.csv'):
    df = populate(input_file_path)
    newDf = pd.DataFrame() 
    df_b = df.drop(df[df.Qualifiers.str[0] != 'B'].index)
    df_s = df.drop(df[df.Qualifiers.str[0] != 'S'].index)
    summed = df.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    summed_b = df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    summed_s = df_s.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).sum(numeric_only=True)
    newDf['DollarVolumeTraded'] = summed.iloc[:,2]
    cats = ['b','s']
    for i in cats:
        if i=='b':
            newDf['DollarVolumeTraded_<'+i+'>'] = summed_b.iloc[:,2]
        elif i=='s':
            newDf['DollarVolumeTraded_<'+i+'>'] = summed_s.iloc[:,2]
    newDf['ShareVolumeTraded'] = summed.iloc[:,1]
    for i in cats:
        if i=='b':
            newDf['ShareVolumeTraded_<'+i+'>'] = summed_b.iloc[:,1]
        elif i=='s':
            newDf['ShareVolumeTraded_<'+i+'>'] = summed_s.iloc[:,1]
    newDf['VWAP'] = newDf.DollarVolumeTraded / newDf.ShareVolumeTraded
    for i in cats:
        if i=='b':
            newDf['VWAP_<'+i+'>'] = newDf['DollarVolumeTraded_<b>'] / newDf['ShareVolumeTraded_<b>']
        elif i=='s':
            newDf['VWAP_<'+i+'>'] = newDf['DollarVolumeTraded_<s>'] / newDf['ShareVolumeTraded_<s>']  
    newDf['LogReturn'] = np.log(newDf.VWAP) - np.log(newDf.VWAP.shift(1)) 
    newDf = newDf.fillna('Not Applicable')
    newDf.to_csv(output_file_path)
    return newDf    

def tradeCount(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_tradeCount.csv'):            
    df = populate(input_file_path)
    newDf = pd.DataFrame() 
    newDf['TradeCount'] = df.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).size()
    df_b = df.drop(df[df.Qualifiers.str[0] != 'B'].index)
    df_s = df.drop(df[df.Qualifiers.str[0] != 'S'].index)
    cats = ['b','s']
    for i in cats:
        if i=='b':
            newDf['TradeCount_<'+i+'>'] = df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).size()
        elif i=='s':
            newDf['TradeCount_<'+i+'>'] = df_s.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).size()
    newDf = newDf.fillna('Not Applicable')
    newDf.to_csv(output_file_path)
    return newDf

def effectiveSpread(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_effectiveSpread.csv'):
    df = populate(input_file_path)
    newDf = pd.DataFrame() 
    cats = ['b','s']
    df_b = df.drop(df[df.Qualifiers.str[0] != 'B'].index)
    df_s = df.drop(df[df.Qualifiers.str[0] != 'S'].index)
    exchanges = df['#RIC'].str.split('.',expand=True)[1].unique().tolist() 
    for m in exchanges:
        df_m = df.drop(df[df['#RIC'].str.split('.',expand=True)[1] != m].index)
        df_m = df_m.set_index('Date-Time')
        df_m = df_m[['Bid Price','Ask Price']]
        averaged_m = df_m.groupby(pd.Grouper(level='Date-Time', axis=0, freq=(str(period)+'T'))).mean(numeric_only=True)
        newDf['QuotedSpread_<'+m+'>'] = averaged_m.iloc[:,1] - averaged_m.iloc[:,0]
        newDf['PercentageSpread_<'+m+'>'] = (averaged_m.iloc[:,1] - averaged_m.iloc[:,0])/((averaged_m.iloc[:,1] + averaged_m.iloc[:,0])/2)
        df_v = df.drop(df[df['#RIC'].str.split('.', expand=True)[1] != m].index)
        df_v = df_v.set_index('Date-Time')
        df_v = df_v[['Bid Size','Ask Size']]
        max_v = df_v.groupby(pd.Grouper(level='Date-Time', axis=0, freq=(str(period)+'T'))).max(numeric_only=True)
        newDf['QuotedDollarDepth_<'+m+'>'] = ((max_v.iloc[:,1] * averaged_m.iloc[:,0]) + (averaged_m.iloc[:,1] * max_v.iloc[:,0]))/2
        newDf['QuotedShareDepth_<'+m+'>'] = (max_v.iloc[:,1] + max_v.iloc[:,0])/2
        for i in cats:
            if i=='b':
                temp = df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).nth(-1)
                newDf['EffectiveSpread_<k,'+m+','+i+',b>'] = np.log(temp.Price) - np.log(df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).mean(numeric_only=True).iloc[:,0])
            elif i=='s':
                temp = df_s.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).nth(-1)
                newDf['EffectiveSpread_<k,'+m+','+i+',b>'] = np.log(temp.Price) - np.log(df_s.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).mean(numeric_only=True).iloc[:,0])
    newDf = newDf.fillna('Not Applicable')
    newDf.to_csv(output_file_path)
    return newDf
    
def realisedSpread(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_realisedSpread.csv'):
    df = populate(input_file_path)
    newDf = pd.DataFrame() 
    cats = ['b','s']
    df_b = df.drop(df[df.Qualifiers.str[0] != 'B'].index)
    df_s = df.drop(df[df.Qualifiers.str[0] != 'S'].index)
    exchanges = df['#RIC'].str.split('.',expand=True)[1].unique().tolist() 
    for m in exchanges:
        df_m = df.drop(df[df['#RIC'].str.split('.',expand=True)[1] != m].index)
        df_m = df_m.set_index('Date-Time')
        df_m = df_m[['Bid Price','Ask Price']]
        averaged_m = df_m.groupby(pd.Grouper(level='Date-Time', axis=0, freq=(str(period)+'T'))).mean(numeric_only=True)
        newDf['QuotedSpread_<'+m+'>'] = averaged_m.iloc[:,1] - averaged_m.iloc[:,0]
        newDf['PercentageSpread_<'+m+'>'] = (averaged_m.iloc[:,1] - averaged_m.iloc[:,0])/((averaged_m.iloc[:,1] + averaged_m.iloc[:,0])/2)
        df_v = df.drop(df[df['#RIC'].str.split('.', expand=True)[1] != m].index)
        df_v = df_v.set_index('Date-Time')
        df_v = df_v[['Bid Size','Ask Size']]
        max_v = df_v.groupby(pd.Grouper(level='Date-Time', axis=0, freq=(str(period)+'T'))).max(numeric_only=True)
        newDf['QuotedDollarDepth_<'+m+'>'] = ((max_v.iloc[:,1] * averaged_m.iloc[:,0]) + (averaged_m.iloc[:,1] * max_v.iloc[:,0]))/2
        newDf['QuotedShareDepth_<'+m+'>'] = (max_v.iloc[:,1] + max_v.iloc[:,0])/2
        for i in cats:
            if i=='b':
                temp = df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).nth(-1)
                newDf['RealisedSpread_<k,'+m+','+i+',b>'] = 2 * (np.log(temp.Price) - np.log(df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period+period)+'T'))).mean(numeric_only=True).iloc[:,0]))
            elif i=='s':
                temp = df_s.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).nth(-1)
                newDf['RealisedSpread_<k,'+m+','+i+',b>'] = 2 * (np.log(df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period+period)+'T'))).mean(numeric_only=True).iloc[:,0]) - np.log(temp.Price))
    newDf = newDf.fillna('Not Applicable')
    newDf.to_csv(output_file_path)
    return newDf
    
def priceImpact(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_priceImpact.csv'):
    df = populate(input_file_path)
    newDf = pd.DataFrame() 
    cats = ['b','s']
    df_b = df.drop(df[df.Qualifiers.str[0] != 'B'].index)
    df_s = df.drop(df[df.Qualifiers.str[0] != 'S'].index)
    exchanges = df['#RIC'].str.split('.',expand=True)[1].unique().tolist()  
    for m in exchanges:
        df_m = df.drop(df[df['#RIC'].str.split('.',expand=True)[1] != m].index)
        df_m = df_m.set_index('Date-Time')
        df_m = df_m[['Bid Price','Ask Price']]
        averaged_m = df_m.groupby(pd.Grouper(level='Date-Time', axis=0, freq=(str(period)+'T'))).mean(numeric_only=True)
        newDf['QuotedSpread_<'+m+'>'] = averaged_m.iloc[:,1] - averaged_m.iloc[:,0]
        newDf['PercentageSpread_<'+m+'>'] = (averaged_m.iloc[:,1] - averaged_m.iloc[:,0])/((averaged_m.iloc[:,1] + averaged_m.iloc[:,0])/2)
        df_v = df.drop(df[df['#RIC'].str.split('.', expand=True)[1] != m].index)
        df_v = df_v.set_index('Date-Time')
        df_v = df_v[['Bid Size','Ask Size']]
        max_v = df_v.groupby(pd.Grouper(level='Date-Time', axis=0, freq=(str(period)+'T'))).max(numeric_only=True)
        newDf['QuotedDollarDepth_<'+m+'>'] = ((max_v.iloc[:,1] * averaged_m.iloc[:,0]) + (averaged_m.iloc[:,1] * max_v.iloc[:,0]))/2
        newDf['QuotedShareDepth_<'+m+'>'] = (max_v.iloc[:,1] + max_v.iloc[:,0])/2
        for i in cats:
            if i=='b':
                temp = df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period+period)+'T'))).nth(-1)
                newDf['PriceImpact_<k,'+m+','+i+',b>'] = 2 * (np.log(temp.Price) - np.log(df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).mean(numeric_only=True).iloc[:,0]))
            elif i=='s':
                temp = df_s.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).nth(-1)
                newDf['PriceImpact_<k,'+m+','+i+',b>'] = 2 * (np.log(temp.Price) - np.log(df_b.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period+period)+'T'))).mean(numeric_only=True).iloc[:,0]))
    newDf = newDf.fillna('Not Applicable')
    newDf.to_csv(output_file_path)
    return newDf

def averagePrice(input_file_path:str='dataspace/input.csv', period:int=60, output_file_path:str= 'dataspace/dataspace_averagePrice.csv'):    
    df = populate(input_file_path)
    newDf = pd.DataFrame() 
    averaged = df.groupby(pd.Grouper(key='Date-Time', axis=0, freq=(str(period)+'T'))).mean(numeric_only=True)
    newDf['AveragePrice'] = averaged.Price
    newDf = newDf.fillna('Not Applicable')
    newDf.to_csv(output_file_path)
    return newDf
    
def main():
    dollarVolumeTraded('dataspace/BHPAX_20190717_Allday.csv')
    shareVolumeTraded('dataspace/BHPAX_20190717_Allday.csv')
    vWAP('dataspace/BHPAX_20190717_Allday.csv')
    arithmeticReturn('dataspace/BHPAX_20190717_Allday.csv')
    logReturn('dataspace/BHPAX_20190717_Allday.csv')
    tradeCount('dataspace/BHPAX_20190717_Allday.csv')
    effectiveSpread('dataspace/BHPAX_20190717_Allday.csv')
    realisedSpread('dataspace/BHPAX_20190717_Allday.csv')
    priceImpact('dataspace/BHPAX_20190717_Allday.csv')
    averagePrice('dataspace/BHPAX_20190717_Allday.csv')

if __name__ == "__main__":
    main()
# regression_pipeline.py

import pandas as pd
import joblib
from sklearn.datasets import fetch_california_housing
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_squared_error, r2_score
from typing import Union

def fetch_sklearn_dataset(
    dataset_name: str = 'california_housing',
    output_csv_file: str = 'dataspace/housing_dataset.csv'
):
    """
    Fetches a dataset from scikit-learn's datasets and saves it as a CSV file.

    Parameters:
    - dataset_name (str): Name of the dataset to fetch ('california_housing').
    - output_csv_file (str): Path to save the dataset as a CSV file.

    Returns:
    - pd.DataFrame: The dataset as a pandas DataFrame.
    """
    if dataset_name == 'california_housing':
        data = fetch_california_housing(as_frame=True)
    else:
        raise ValueError("Unsupported dataset_name. Choose 'california_housing'.")

    df = pd.concat([data.data, data.target.rename('target')], axis=1)
    df.to_csv(output_csv_file, index=False)
    return df

def train_sklearn_regression_model(
    input_file_path: str = 'dataspace/housing_dataset.csv',
    target_column: str = 'target',
    test_size: float = 0.2,
    random_state: int = 42,
    model_output_path: str = 'dataspace/sklearn_regression_model.pkl',
    performance_output_path: str = 'dataspace/sklearn_regression_performance.txt'
):
    """
    Trains a Linear Regression model and saves it.

    Parameters:
    - input_file_path (str): Path to the dataset CSV file.
    - target_column (str): Name of the target column.
    - test_size (float): Fraction of data to use as test set.
    - random_state (int): Seed for reproducibility.
    - model_output_path (str): Path to save the trained model.
    - performance_output_path (str): Path to save performance metrics.

    Returns:
    - dict: Dictionary containing performance metrics.
    """
    df = pd.read_csv(input_file_path)
    X = df.drop(columns=[target_column])
    y = df[target_column]

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=test_size, random_state=random_state
    )

    model = LinearRegression()
    model.fit(X_train, y_train)

    y_pred = model.predict(X_test)
    mse = mean_squared_error(y_test, y_pred)
    r2 = r2_score(y_test, y_pred)

    # Save the model
    joblib.dump(model, model_output_path)

    # Save performance metrics
    with open(performance_output_path, 'w') as f:
        f.write(f'MSE: {mse}\n')
        f.write(f'R2 Score: {r2}\n')

    return {'MSE': mse, 'R2 Score': r2}

def sklearn_model_predict(
    model_path: str = 'dataspace/sklearn_regression_model.pkl',
    input_data_path: str = 'dataspace/housing_test.csv',
    output_data_path: str = 'dataspace/housing_predictions.csv'
):
    """
    Loads a trained sklearn model and makes predictions on new data.

    Parameters:
    - model_path (str): Path to the saved sklearn model.
    - input_data_path (str): Path to the CSV file containing new data.
    - output_data_path (str): Path to save the predictions.

    Returns:
    - pd.Series: Predictions made by the model.
    """
    model = joblib.load(model_path)
    X_new = pd.read_csv(input_data_path)
    predictions = model.predict(X_new)

    # Save predictions to CSV
    output_df = X_new.copy()
    output_df['Predicted_Target'] = predictions
    output_df.to_csv(output_data_path, index=False)

    return predictions

# sentiment_analysis_pipeline.py

import pandas as pd
import nltk
import joblib
from nltk.corpus import movie_reviews
from sklearn.model_selection import train_test_split
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, classification_report

def fetch_nltk_movie_reviews(output_csv_file: str = 'dataspace/movie_reviews.csv'):
    """
    Fetches the NLTK movie reviews dataset and saves it as a CSV file.

    Parameters:
    - output_csv_file (str): Path to save the dataset.

    Returns:
    - pd.DataFrame: DataFrame containing reviews and labels.
    """
    nltk.download('movie_reviews', quiet=True)

    documents = [
        (' '.join(movie_reviews.words(fileid)), category)
        for category in movie_reviews.categories()
        for fileid in movie_reviews.fileids(category)
    ]

    reviews = [doc for doc, _ in documents]
    labels = [label for _, label in documents]

    df = pd.DataFrame({'Review': reviews, 'Label': labels})
    df.to_csv(output_csv_file, index=False)
    return df

def train_text_classification_model(
    input_file_path: str = 'dataspace/movie_reviews.csv',
    text_column: str = 'Review',
    target_column: str = 'Label',
    test_size: float = 0.2,
    random_state: int = 42,
    model_output_path: str = 'dataspace/text_classification_model.pkl',
    performance_output_path: str = 'dataspace/text_classification_performance.txt'
):
    """
    Trains a text classification model using Logistic Regression.

    Parameters:
    - input_file_path (str): Path to the dataset CSV file.
    - text_column (str): Name of the text column.
    - target_column (str): Name of the target column.
    - test_size (float): Fraction of data to use as test set.
    - random_state (int): Seed for reproducibility.
    - model_output_path (str): Path to save the trained model.
    - performance_output_path (str): Path to save performance metrics.

    Returns:
    - dict: Dictionary containing performance metrics.
    """
    df = pd.read_csv(input_file_path)
    X = df[text_column]
    y = df[target_column]

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=test_size, random_state=random_state
    )

    vectorizer = TfidfVectorizer(stop_words='english', max_features=5000)
    X_train_vect = vectorizer.fit_transform(X_train)
    X_test_vect = vectorizer.transform(X_test)

    model = LogisticRegression(max_iter=1000)
    model.fit(X_train_vect, y_train)

    y_pred = model.predict(X_test_vect)
    accuracy = accuracy_score(y_test, y_pred)
    report = classification_report(y_test, y_pred)

    # Save the model and vectorizer
    joblib.dump((model, vectorizer), model_output_path)

    # Save performance metrics
    with open(performance_output_path, 'w') as f:
        f.write(f'Accuracy: {accuracy}\n')
        f.write(f'Classification Report:\n{report}')

    return {'Accuracy': accuracy, 'Report': report}

def text_model_predict(
    model_path: str = 'dataspace/text_classification_model.pkl',
    input_texts: list = [
        "I absolutely loved this movie. The story was gripping and the characters were well-developed.",
        "This was the worst film I have ever seen. Completely a waste of time."
    ],
    output_file_path: str = 'dataspace/text_predictions.csv'
):
    """
    Uses the trained text classification model to predict sentiments of new texts.

    Parameters:
    - model_path (str): Path to the saved model and vectorizer.
    - input_texts (list): List of texts to classify.
    - output_file_path (str): Path to save the predictions.

    Returns:
    - pd.DataFrame: DataFrame containing texts and their predicted labels.
    """
    model, vectorizer = joblib.load(model_path)
    X_new_vect = vectorizer.transform(input_texts)
    predictions = model.predict(X_new_vect)

    df = pd.DataFrame({'Text': input_texts, 'Predicted_Label': predictions})
    df.to_csv(output_file_path, index=False)
    return df
import pandas as pd

def convert_to_long_format(input_file: str = "dataspace/World_Development_Indicators.csv", output_file: str = "dataspace/world_dev_indicator_long.csv", id_vars: list = ['Country Name', 'Country Code', 'Series Name', 'Series Code'], key_name: str = "Year", value_name: str = "Money"):
    df = pd.read_csv(input_file)
    df_long = pd.melt(df, id_vars=id_vars, var_name=key_name, value_name=value_name)
    df_long.to_csv(output_file, index=False)
    return df_long

def convert_to_integer(input_file: str = "dataspace/world_dev_indicator_long.csv", output_file: str = "dataspace/world_dev_indicator_long.csv", columns: list = ['Year']):
    df = pd.read_csv(input_file)
    for column in columns:
        df[column] = df[column].astype(int)
    df.to_csv(output_file, index=False)
    return df

def filter_by_multiple_val(input_file: str = "dataspace/world_dev_indicator_long.csv", output_file: str = "dataspace/gdp_series_all_2019.csv", columns: list = ['Series Name', 'Year'], values: list = ['GDP (current US$)', 2019]):
    df = pd.read_csv(input_file)
    for column, value in zip(columns, values):
        df = df[df[column] == value]
    df.to_csv(output_file, index=False)
    return df

def remove_row_with_na_specified_columns(input_file: str = "dataspace/gdp_series_all_2019.csv", output_file: str = "dataspace/gdp_series_all_2019.csv", columns: list = ['Money']):
    df = pd.read_csv(input_file)
    for column in columns:
        df = df.dropna(subset=[column])
    df.to_csv(output_file, index=False)
    return df

def get_top_n_rows(input_file: str = "dataspace/gdp_series_all_2019.csv", output_file: str = "dataspace/gdp_series_top_5_2019.csv", n: int = 5):
    df = pd.read_csv(input_file)
    df_top_n = df.head(n)
    df_top_n.to_csv(output_file, index=False)
    return df_top_n

def get_bottom_n_rows(input_file: str = "dataspace/gdp_series_all_2019.csv", output_file: str = "dataspace/gdp_series_bottom_5_2019.csv", n: int = 5):
    df = pd.read_csv(input_file)
    df_bottom_n = df.tail(n)
    df_bottom_n.to_csv(output_file, index=False)
    return df_bottom_n

def concatenate_dataframes(input_files: list = ['dataspace/gdp_series_top_5_2019.csv', 'dataspace/gdp_series_bottom_5_2019.csv'], output_file: str = "dataspace/gdp_series_2019.csv"):
    df_list = [pd.read_csv(file) for file in input_files]
    df_concatenated = pd.concat(df_list)
    df_concatenated.to_csv(output_file, index=False)
    return df_concatenated

def filter_by_val_list(input_file: str = "dataspace/world_dev_indicator_long.csv", output_file: str = "dataspace/gdp_series_all.csv", column: str = 'Country Name', value_path: str = "dataspace/gdp_series_2019.csv"):
    df = pd.read_csv(input_file)
    value_df = pd.read_csv(value_path)
    value_list = value_df[column].tolist()
    df = df[df[column].isin(value_list)]
    df.to_csv(output_file, index=False)
    return df

def filter_by_single_val(input_file: str = "dataspace/gdp_series_all.csv", output_file: str = "dataspace/gdp_series_all.csv", column: str = 'Series Name', value: str = 'GDP (current US$)'):
    df = pd.read_csv(input_file)
    df = df[df[column] == value]
    df.to_csv(output_file, index=False)
    return df
# clustering_pipeline.py

import pandas as pd
from sklearn.datasets import load_iris
from sklearn.cluster import KMeans
import plotly.express as px

def fetch_iris_dataset(output_csv_file: str = 'dataspace/iris_dataset.csv'):
    """
    Fetches the Iris dataset and saves it as a CSV file.

    Parameters:
    - output_csv_file (str): Path to save the dataset as a CSV file.

    Returns:
    - pd.DataFrame: The dataset as a pandas DataFrame.
    """
    data = load_iris(as_frame=True)
    df = pd.concat([data.data, data.target.rename('target')], axis=1)
    df.to_csv(output_csv_file, index=False)
    return df

def perform_kmeans_clustering(
    input_file_path: str = 'dataspace/iris_dataset.csv',
    n_clusters: int = 3,
    output_file_path: str = 'dataspace/iris_clusters.csv'
):
    """
    Performs KMeans clustering on the dataset.

    Parameters:
    - input_file_path (str): Path to the dataset CSV file.
    - n_clusters (int): Number of clusters.
    - output_file_path (str): Path to save the clustered data.

    Returns:
    - pd.DataFrame: DataFrame with cluster labels.
    """
    df = pd.read_csv(input_file_path)
    X = df.drop(columns=['target'])

    kmeans = KMeans(n_clusters=n_clusters, random_state=42)
    df['Cluster'] = kmeans.fit_predict(X)

    df.to_csv(output_file_path, index=False)
    return df

def plotly_cluster_visualization(
    input_csv_file: str = 'dataspace/iris_clusters.csv',
    x_axis: str = 'sepal length (cm)',
    y_axis: str = 'sepal width (cm)',
    cluster_col: str = 'Cluster',
    output_html_file: str = 'dataspace/iris_clusters_plot.html'
):
    """
    Visualizes clusters using Plotly.

    Parameters:
    - input_csv_file (str): Path to the CSV file with clustered data.
    - x_axis (str): Column name for the X-axis.
    - y_axis (str): Column name for the Y-axis.
    - cluster_col (str): Column name for cluster labels.
    - output_html_file (str): Path to save the HTML plot.

    Returns:
    - None
    """
    df = pd.read_csv(input_csv_file)
    fig = px.scatter(df, x=x_axis, y=y_axis, color=cluster_col)
    fig.write_html(output_html_file)
