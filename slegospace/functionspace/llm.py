import json
import openai
import os
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import pytesseract
import csv
import json
import openai
import os
import PyPDF2
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import pytesseract
import csv
from openai import OpenAI

def chatbot_huggingface_api(API_URL: str = "https://api-inference.huggingface.co/models/mistralai/Mistral-7B-Instruct-v0.2",
                            api_key: str = "your_api_key",
                            user_input_file: str = '',
                            output_text_file: str = 'dataspace/gpt_output_text.txt',
                            output_json_file: str = 'dataspace/gpt_output_full.json',
                            query: dict = {"inputs": "Can you please let us know more details about your "}):
    """
    Sends a query, optionally augmented with contents from various file types, to the specified Hugging Face API endpoint.
    
    This function supports processing inputs from text, PDF, DOCX, XLSX, image files (for OCR), JSON, and CSV files. The
    contents of the file are appended to a base query provided in the `query` parameter. If a file is specified but does
    not exist, the function will raise a FileNotFoundError. Unsupported file types will raise a ValueError.

    Parameters:
        API_URL (str): The URL of the Hugging Face API model to which the request will be sent.
        api_key (str): The API key required for authentication with the Hugging Face API.
        user_input_file (str): The path to the file whose contents are to be appended to the query. If empty, only the base query is used.
        output_text_file (str): The path where the plain text part of the response will be saved.
        output_json_file (str): The path where the full JSON response from the API will be saved.
        query (dict): A dictionary containing the base query. Expected to have at least a key 'inputs' with a starting string.

    Returns:
        dict: A dictionary containing the JSON response from the API. If there's an HTTP error, it returns a dictionary
              containing the error message and status code.

    Raises:
        FileNotFoundError: If the specified `user_input_file` does not exist.
        ValueError: If the file extension of `user_input_file` is not supported.

    Example Usage:
        result = chatbot_huggingface_api(
            api_key="your_api_key_here",
            user_input_file="path/to/input.txt",
            query={"inputs": "Please analyze the following data: "}
        )
        print(result)
    """
    # Function implementation here
    # Initialize the combined message with the query's input
    combined_message = query["inputs"]

    # Process file if specified and exists
    if user_input_file and os.path.exists(user_input_file):
        file_extension = user_input_file.split('.')[-1].lower()

        if file_extension == 'txt':
            with open(user_input_file, 'r') as file:
                file_contents = file.read().strip()
            combined_message += f"\n\n==== Text File Input ====\n\n{file_contents}"
        elif file_extension == 'pdf':
            with open(user_input_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pdf_contents = ' '.join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            combined_message += f"\n\n==== PDF File Input ====\n\n{pdf_contents}"
        elif file_extension == 'docx':
            doc = Document(user_input_file)
            docx_contents = ' '.join([para.text for para in doc.paragraphs])
            combined_message += f"\n\n==== DOCX File Input ====\n\n{docx_contents}"
        elif file_extension == 'xlsx':
            workbook = load_workbook(filename=user_input_file)
            sheet = workbook.active
            xlsx_contents = ' '.join([str(cell.value) for row in sheet for cell in row if cell.value is not None])
            combined_message += f"\n\n==== XLSX File Input ====\n\n{xlsx_contents}"
        elif file_extension in ['png', 'jpg', 'jpeg']:
            img = Image.open(user_input_file)
            image_text = pytesseract.image_to_string(img)
            combined_message += f"\n\n==== Image File Input (OCR) ====\n\n{image_text}"
        elif file_extension == 'json':
            with open(user_input_file, 'r') as file:
                json_data = json.load(file)
                json_contents = json.dumps(json_data, indent=4)
            combined_message += f"\n\n==== JSON File Input ====\n\n{json_contents}"
        elif file_extension == 'csv':
            with open(user_input_file, mode='r', newline='', encoding='utf-8') as file:
                reader = csv.reader(file)
                csv_contents = ' '.join([','.join(row) for row in reader])
            combined_message += f"\n\n==== CSV File Input ====\n\n{csv_contents}"
        else:
            raise ValueError("Unsupported file extension")
    elif user_input_file and not os.path.exists(user_input_file):
        raise FileNotFoundError("The specified input file does not exist or is not accessible")

    # Send request to the Hugging Face API
    headers = {"Authorization": f"Bearer {api_key}"}
    response = requests.post(API_URL, headers=headers, json={"inputs": combined_message})
    if response.status_code == 200:
        response_data = response.json()
        # Save response data to JSON file
        with open(output_json_file, 'w') as jsonfile:
            json.dump(response_data, jsonfile, indent=4)
        # Extract text part of response and save to text file
        if 'generated_text' in response_data:
            with open(output_text_file, 'w') as textfile:
                textfile.write(response_data['generated_text'])
        return response_data
    else:
        return {'error': 'Failed to get a valid response', 'status_code': response.status_code}


def chatgpt_chat(model:str='gpt-3.5-turbo',
                  user_input_file:str='dataspace/user_text_input.txt',
                  output_text_file:str='dataspace/gpt_output_text.txt',
                  output_json_file:str='dataspace/gpt_output_full.json',
                  temperature:int=1, 
                  max_tokens:int=256, 
                  top_p:int=1, 
                  frequency_penalty:int=0, 
                  presence_penalty:int=0,
                  api_key:str='sk-CiO5GzpXbxZQsMuKEQEkT3BlbkFJz4LS3FuI3f5NqmF1BXO', 
                  user_message:str='Summarize:',):

    '''
    This function interfaces with OpenAI's GPT model to process a text input and obtain a generated response.
    It reads an additional input from a file, merges it with the user's direct input, and sends the combined
    content to the API. The response is then saved to both text and JSON files.

    Parameters:
        api_key (str): Your OpenAI API key.
        model (str): Identifier for the model version to use.
        user_message (str): Direct user input as a string.
        user_input_file (str): Path to a file containing additional text to send to the API.
        output_file (str): Path to save the plain text response from the API.
        output_json_file (str): Path to save the full JSON response from the API.
        temperature (float): Controls randomness in response generation. Higher is more random.
        max_tokens (int): Maximum length of the response.
        top_p (float): Controls diversity via nucleus sampling: 0.1 means only top 10% probabilities considered.
        frequency_penalty (float): Decreases likelihood of repeating words.
        presence_penalty (float): Decreases likelihood of repeating topics.

    Returns:
        str: The text response from the API.
    '''

    # Initialize variables to store responses
    ans = None
    ans_dict = {}

    combined_message = user_message  # Start with the user's direct message

    # Determine the file type and read content accordingly
    if os.path.exists(user_input_file):
        file_extension = user_input_file.split('.')[-1].lower()
        if file_extension == 'txt':
            with open(user_input_file, 'r') as file:
                file_contents = file.read().strip()
            combined_message += f"\n\n==== Text File Input ====\n\n{file_contents}"
        elif file_extension == 'pdf':
            with open(user_input_file, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pdf_contents = ' '.join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            combined_message += f"\n\n==== PDF File Input ====\n\n{pdf_contents}"
        elif file_extension == 'docx':
            doc = Document(user_input_file)
            docx_contents = ' '.join([para.text for para in doc.paragraphs])
            combined_message += f"\n\n==== DOCX File Input ====\n\n{docx_contents}"
        elif file_extension == 'xlsx':
            workbook = load_workbook(filename=user_input_file)
            sheet = workbook.active
            xlsx_contents = ' '.join([str(cell.value) for row in sheet for cell in row if cell.value is not None])
            combined_message += f"\n\n==== XLSX File Input ====\n\n{xlsx_contents}"
        elif file_extension in ['png', 'jpg', 'jpeg']:
            img = Image.open(user_input_file)
            image_text = pytesseract.image_to_string(img)
            combined_message += f"\n\n==== Image File Input (OCR) ====\n\n{image_text}"
        elif file_extension == 'json':
            with open(user_input_file, 'r') as file:
                json_data = json.load(file)
                json_contents = json.dumps(json_data, indent=4)
            combined_message += f"\n\n==== JSON File Input ====\n\n{json_contents}"
        elif file_extension == 'csv':
            with open(user_input_file, mode='r', newline='', encoding='utf-8') as file:
                reader = csv.reader(file)
                csv_contents = ' '.join([','.join(row) for row in reader])
            combined_message += f"\n\n==== CSV File Input ====\n\n{csv_contents}"


    try:
        # Set the API key (consider using environment variables for security)
        client = OpenAI(api_key=api_key)
        # Create a chat completion request with the specified parameters
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": combined_message}],
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=top_p,
            frequency_penalty=frequency_penalty,
            presence_penalty=presence_penalty,
        )

    # response = client.chat.completions.create(
    #     model="gpt-4o",
    #     response_format={ "type": "json_object" },
    #     messages=[
    #         {"role": "system", "content": system_message},
    #         {"role": "user", "content": user_message}
    #     ],
    #     temperature=1,
    #     max_tokens=1280,
    #     top_p=1,
    #     frequency_penalty=0,
    #     presence_penalty=0,
    #     )
    # response_text = response.choices[0].message.content.strip() # response['choices'][0]['message']['content'].strip()
    # return response_text
        # Extract and process the response
        ans_dict = response.to_dict()
        if 'choices' in ans_dict and len(ans_dict['choices']) > 0:
            if 'message' in ans_dict['choices'][0]:
                ans = ans_dict['choices'][0]['message']['content']

        # Save the text response and the full JSON response
        if ans:
            with open(output_text_file, 'w') as f:
                f.write(ans)
        with open(output_json_file, 'w') as json_file:
            json.dump(ans_dict, json_file, indent=4)

    except Exception as e:
        print(f"An unexpected error occurred: {e}")

    return ans
